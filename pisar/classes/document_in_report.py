import datetime
import os
import string

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Mm
from docx.shared import Pt
from pytrovich.enums import NamePart, Gender, Case

from classes.paragraph_settings import ParagraphSettings
from document_prototype import DocumentPrototype
from helpers.text_helper import get_word_declension, get_words_declension, get_month_string, replace_with_glue, \
	glue_number_string

MODEL_PERSONNEL_PATH = "personnel_path"
MODEL_PERSONNEL_DETAILS_PATH = "personnel_details_path"
MODEL_OUTPUT_FOLDER = "output_folder"
MODEL_PERSONS = "persons"
MODEL_MORPHOLOGY = "morphology"
MODEL_JSON_OBJECT = "json_settings"
MODEL_MORPHOLOGY_FOR_NAMES = "petrovich"
MODEL_IS_VALID = "is_valid"
MODEL_CURRENT_SOLDIER = "current_soldier"


class DocumentInReport(DocumentPrototype):
	def __init__(self, data_model):
		super().__init__(data_model)
		self.subfolder_name = None
		self.pages = []
		self.word_document = Document()
		style = self.word_document.styles['Normal']
		font = style.font
		font.name = 'Times New Roman'
		font.size = Pt(14)
		# predefined styles
		self.bold_center_settings = ParagraphSettings()
		self.bold_center_settings.is_bold = True
		self.bold_center_settings.align_center = True

		self.align_left_settings = ParagraphSettings()
		self.align_left_settings.align_left = True

		self.align_right_settings = ParagraphSettings()
		self.align_right_settings.align_right = True

		self.bold_right_settings = ParagraphSettings()
		self.bold_right_settings.align_right = True
		self.bold_right_settings.is_bold = True

		self.align_center_settings = ParagraphSettings()
		self.align_center_settings.align_center = True

		self.align_justify_settings = ParagraphSettings()
		self.align_justify_settings.align_justify = True

		self.bold_justify_settings = ParagraphSettings()
		self.bold_justify_settings.align_justify = True
		self.bold_justify_settings.is_bold = True

		self.ident_align_justify_settings = ParagraphSettings()
		self.ident_align_justify_settings.align_justify = True
		self.ident_align_justify_settings.first_line_indent = Mm(12.5)

		self.bold_title = ParagraphSettings()
		self.bold_title.font_size = Pt(16)
		self.bold_title.is_bold = True
		self.bold_title.align_center = True

		self.underline_settings = ParagraphSettings()
		self.underline_settings.is_underline = True

		self.align_justify_underline = ParagraphSettings()
		self.align_justify_underline.align_justify = True
		self.align_justify_underline.is_underline = True

		self.align_center_underline = ParagraphSettings()
		self.align_center_underline.align_center = True
		self.align_center_underline.is_underline = True

	# self.personnel_info = None
	# if self.data_model is not None and self.data_model[MODEL_PERSONNEL_PATH] is not None:
	# self.personnel_info = PersonnelStorage(self.data_model[MODEL_PERSONNEL_PATH])

	# creates MS Word document
	def render(self, custom_margins=None):
		self.apply_default_settings(custom_margins)
		if self.data_model is not None and self.data_model[MODEL_OUTPUT_FOLDER]:
			full_path_folder = self.data_model[MODEL_OUTPUT_FOLDER]
			# store files in a sub folder with the soldier name
			# or use the dedicated sub folder if it is set
			if self.subfolder_name is not None:
				full_path_folder = os.path.join(full_path_folder, self.subfolder_name)
			else:
				s_info = self.get_soldier_info()
				if s_info is not None:
					full_path_folder = os.path.join(full_path_folder, s_info.full_name)

			if not os.path.exists(full_path_folder):
				os.makedirs(full_path_folder)
			full_path = os.path.join(full_path_folder, self.get_name_for_file())
			self.word_document.save(full_path)
			print(f"Создан документ {full_path}.")

	def add_paragraph(self, text, paragraph_settings):
		if paragraph_settings is None:
			paragraph_settings = ParagraphSettings()
		if paragraph_settings.has_markup():
			p = self.word_document.add_paragraph()
			runner = p.add_run(text)
			if paragraph_settings.is_bold:
				runner.bold = True
			if paragraph_settings.is_italic:
				runner.italic = True
			if paragraph_settings.is_underline:
				runner.underline = True
			if paragraph_settings.font_size > 0:
				runner.font.size = paragraph_settings.font_size
		else:
			p = self.word_document.add_paragraph(text)

		# TODO use style
		pf = p.paragraph_format
		pf.left_indent = Pt(0)
		pf.right_indent = Pt(0)
		pf.space_before = Pt(0)
		pf.space_after = Pt(0)

		if paragraph_settings.align_left:
			p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
		if paragraph_settings.align_center:
			p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
		if paragraph_settings.align_right:
			p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
		if paragraph_settings.align_justify:
			p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

		if paragraph_settings.first_line_indent > 0:
			pf.first_line_indent = paragraph_settings.first_line_indent

		if paragraph_settings.left_indent > 0:
			pf.left_indent = paragraph_settings.left_indent

		if paragraph_settings.right_indent > 0:
			pf.right_indent = paragraph_settings.right_indent

		if paragraph_settings.line_spacing > 0:
			pf.line_spacing = paragraph_settings.line_spacing

		return p

	def add_empty_paragraphs_spacing(self, how_many_rows, line_spacing):
		num_row = 1
		paragraph_settings = ParagraphSettings()
		paragraph_settings.line_spacing = line_spacing
		while num_row <= how_many_rows:
			self.add_paragraph("", paragraph_settings)
			num_row = num_row + 1

	def add_empty_paragraphs(self, how_many_rows):
		num_row = 1
		while num_row <= how_many_rows:
			self.add_paragraph("", None)
			num_row = num_row + 1

	def add_empty_paragraphs_small(self, how_many_rows):
		paragraph_settings = ParagraphSettings()
		paragraph_settings.font_size = Pt(8)
		num_row = 1
		while num_row <= how_many_rows:
			self.add_paragraph(" ", paragraph_settings)
			num_row = num_row + 1

	def apply_default_settings(self, custom_margins=None):
		sections = self.word_document.sections
		top_margin = 20
		bottom_margin = 20
		left_margin = 30
		right_margin = 10
		if custom_margins is not None:
			top_margin = custom_margins.top_margin
			bottom_margin = custom_margins.bottom_margin
			left_margin = custom_margins.left_margin
			right_margin = custom_margins.right_margin

		for section in sections:
			section.top_margin = Mm(top_margin)
			section.bottom_margin = Mm(bottom_margin)  # 15 mm in original document
			section.left_margin = Mm(left_margin)
			section.right_margin = Mm(right_margin)
			# settings for page as A4
			section.page_height = Mm(297)
			section.page_width = Mm(210)
			section.header_distance = Mm(12.7)
			section.footer_distance = Mm(12.7)

	def set_column_width(self, table, index, size):
		for cell in table.columns[index].cells:
			cell.width = Mm(size)

	def add_table(self, captions, rows_data, table_settings=None):
		cols_width = []
		ps = None
		alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
		font_size = None
		if table_settings is not None:
			if "ps" in table_settings:
				ps = table_settings["ps"]
			if "cols_width" in table_settings:
				cols_width = table_settings["cols_width"]
			if "alignment" in table_settings:
				alignment = table_settings["alignment"]
			if "font_size" in table_settings:
				font_size = table_settings["font_size"]

		# if all captions are empty, don't render header
		not_empty = False
		for caption in captions:
			if len(caption) > 0:
				not_empty = True
				break

		rows_count = len(rows_data)
		if not_empty:
			rows_count = rows_count + 1
		table = self.word_document.add_table(rows=rows_count, cols=len(captions))
		table.style = 'Table Grid'
		table.allow_autofit = False
		table.alignment = alignment
		# TODO put table in the center and set margins 0;0
		hdr_cells = table.rows[0].cells
		num_column = 0

		if not_empty:
			for caption in captions:
				cell = hdr_cells[num_column]
				p = cell.paragraphs[0]
				if ps is None:
					p.text = caption
				else:
					# TODO use common method to apply settings
					runner = p.add_run(caption)
					if ps.is_bold:
						runner.bold = True
					if ps.font_size > 0:
						runner.font.size = ps.font_size

				p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
				cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
				num_column = num_column + 1

		# set column width
		if cols_width is not None:
			for col_index in range(0, len(cols_width)):
				self.set_column_width(table, col_index, cols_width[col_index])

		# process rows
		num_row = 1
		if not not_empty:
			num_row = 0
		for row in rows_data:
			cells = table.rows[num_row].cells
			col_ind = 0
			for c_data in row:
				cell = cells[col_ind]
				if font_size is None:
					cell.text = c_data
				else:
					p = cell.paragraphs[0]
					runner = p.add_run(str(c_data))
					runner.font.size = font_size
				col_ind = col_ind + 1
			num_row = num_row + 1

	def add_paragraph_left_right(self, left_text: string, right_text: string, font_size=None, table_alignment=None, is_bold=False):
		table = self.word_document.add_table(rows=1, cols=2)
		if table_alignment is None:
			table_alignment = WD_TABLE_ALIGNMENT.CENTER
		table.alignment = table_alignment
		cells = table.rows[0].cells
		p0 = cells[0].paragraphs[0]
		p0.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
		r_left = p0.add_run(left_text)
		p1 = cells[1].paragraphs[0]
		p1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
		r_right = p1.add_run(right_text)
		if font_size is not None:
			r_left.font.size = font_size
			r_right.font.size = font_size
		if is_bold:
			r_left.bold = True
			r_right.bold = True

		return r_left, r_right

	def get_soldier_info(self):
		return self.data_model[MODEL_CURRENT_SOLDIER]

	def get_report_settings(self):
		return self.data_model[MODEL_JSON_OBJECT]

	def get_person_name_routines(self, full_name, cs):
		maker = self.data_model[MODEL_MORPHOLOGY_FOR_NAMES]
		if maker is None:
			print("Внутренняя ошибка. Морфоанализатор для имён не создан.")
			return ""

		if full_name is None or len(full_name) == 0:
			return ""

		name_tokens = full_name.split(" ")
		surname = name_tokens[0]
		first_name = ""
		if len(name_tokens) > 1:
			first_name = name_tokens[1]
		middle_name = ""
		if len(name_tokens) > 2:
			middle_name = name_tokens[2]

		rs = self.get_report_settings()
		gender = Gender.MALE
		if "gender" in rs and rs["gender"].casefold() == "ж":
			gender = Gender.FEMALE

		sn = maker.make(NamePart.LASTNAME, gender, cs, surname)
		fn = maker.make(NamePart.FIRSTNAME, gender, cs, first_name)
		mn = maker.make(NamePart.MIDDLENAME, gender, cs, middle_name)

		return f"{sn} {fn} {mn}".strip()

	def get_person_name_gent(self, full_name):
		return self.get_person_name_routines(full_name, Case.GENITIVE)

	def get_person_name_instr(self, full_name):
		return self.get_person_name_routines(full_name, Case.INSTRUMENTAL)

	def get_person_name_datv(self, full_name):
		return self.get_person_name_routines(full_name, Case.DATIVE)

	def get_person_name_declension(self, full_name, declension_type):
		if declension_type == 1:
			result = self.get_person_name_gent(full_name)
		else:
			if declension_type == 2:
				result = self.get_person_name_instr(full_name)
			else:
				if declension_type == 3:
					result = self.get_person_name_datv(full_name)
				else:
					result = full_name

		return result

	# format = Петров Алексей Сергеевич -> А. Петров
	def get_person_name_short_format_1(self, full_name):
		name_tokens = full_name.split(" ")
		if len(name_tokens) < 2:
			print(f"Не удалось преобразовать в нужный формат: {full_name}")
			return full_name
		surname = name_tokens[0]
		first_name = name_tokens[1]
		return f"{first_name[0]}. {surname}"

	# format = Петров Алексей Сергеевич -> Петров А.С.
	def get_person_name_short_format_2(self, full_name, declension_type):
		if full_name is None or len(full_name) == 0:
			return full_name
		full_name = self.get_person_name_declension(full_name, declension_type)
		name_tokens = full_name.split(" ")
		surname = name_tokens[0]
		first_name = ""
		second_name = ""
		if len(name_tokens) > 1:
			first_name = name_tokens[1]
		if len(name_tokens) > 2:
			second_name = name_tokens[2]
		result = surname
		if len(first_name) > 0:
			result = result + f" {first_name[0]}."
		if len(second_name) > 0:
			result = result + f" {second_name[0]}."
		return result

	# format = 16-04-2023/16.04.2023 -> 16 апреля 2023 года
	def get_date_format_1(self, date_str):
		tokens = date_str.split(".")
		if len(tokens) != 3:
			return date_str
		# print(f"Не удалось определить формат даты {date_str}")
		d = int(tokens[0])
		m = int(tokens[1])
		y = int(tokens[2])
		return replace_with_glue(f"{d} {get_month_string(m)} {y} года")

	# declension_type. 0 (without), 1 (gent), 2 (ablt), 3 (datv)
	def get_person_full_str(self, settings):
		s_info = self.get_soldier_info()
		sld_position = ""  # if not required
		if settings.militaryman_required:
			sld_position = get_word_declension(self.get_morph(), "военнослужащий", settings.declension_type)
		else:
			if settings.position_required:
				sld_position = get_word_declension(self.get_morph(), s_info.position, settings.declension_type)

		sld_rank = ""
		if settings.rank_required:
			sld_rank = self.get_person_rank(s_info.rank, settings.declension_type)

		# TODO battalion must be a variable
		address = ""
		if settings.address_required:
			address = replace_with_glue("2 стрелкового батальона")
			if settings.military_unit_required:
				address = f"{address} войсковой части {self.get_military_unit()}"
			if not settings.battalion_only:
				address = f"{self.get_soldier_address(1)} {address}"
		# address = f"{s_info.squad} стрелкового отделения {s_info.platoon} стрелкового взвода {s_info.company} стрелковой роты " + address

		dob_str = ""
		if settings.dob_required:
			if settings.is_dob_short:
				dob_str = s_info.get_dob() + " г.р."
			else:
				dob_str = self.get_date_format_1(s_info.get_dob()) + " рождения"

		result = f"{sld_position} {address} {sld_rank}"
		if settings.person_name_required:
			full_name = self.get_person_name_declension(s_info.full_name, settings.declension_type)
			result = result + f" {full_name}"

		if len(dob_str) > 0:
			result = result + ", " + dob_str
		return result.strip()

	def get_person_rank(self, rnk, declension_type):
		if declension_type != 0:
			rnk = get_word_declension(self.get_morph(), rnk, declension_type)
		if self.get_report_settings()["is_guard"]:
			rnk = "гвардии " + rnk
		return rnk

	def get_commander_company(self):
		return self.get_commander_generic("commander_company", "КОМАНДИРА РОТЫ", 0, False)

	def get_commander_company_full_str(self, declension_type, need_capitalize=True):
		text = "[ВСТАВЬТЕ СВЕДЕНИЯ О КОМАНДИРЕ РОТЫ]"
		commander_info = self.get_commander_generic("commander_company", "КОМАНДИРА РОТЫ", declension_type, False)
		if commander_info["found"]:
			m_unit = self.get_military_unit()
			pos = commander_info['position']
			if need_capitalize:
				pos = pos.capitalize()

			text = f"{pos} войсковой части {m_unit} {commander_info['rank']} {commander_info['name']}"
		return text

	def get_morph(self):
		return self.data_model[MODEL_MORPHOLOGY]

	def get_commander_platoon_full_str(self, declension_type):
		text = "[ВСТАВЬТЕ СВЕДЕНИЯ О КОМАНДИРЕ ВЗВОДА]"
		commander_info = self.get_commander_generic("commander_platoon", "КОМАНДИРА ВЗВОДА", declension_type, False)
		if commander_info["found"]:
			# TODO need to use a property from commander_info
			# platoon = get_words_declension(self.get_morph(), self.get_soldier_info().platoon, 1)
			text = f"{commander_info['position']} войсковой части {self.get_military_unit()} {commander_info['rank']} {commander_info['name']}"
		return text

	# TODO refactor this method
	def get_commander_generic_full_str(self,
	                                   settings_key,
	                                   declension_type,
	                                   need_сapitalize=True):
		text = "[ВСТАВЬТЕ СВЕДЕНИЯ О КОМАНДИРЕ]"

		commander_info = self.get_commander_generic(settings_key, text, declension_type, False)
		if commander_info["found"]:
			pos = commander_info['position']
			if need_сapitalize:
				pos = pos.capitalize()
			text = f"{pos} войсковой части {self.get_military_unit()} {commander_info['rank']} {commander_info['name']}"
		return text

	def get_military_unit(self):
		return self.get_report_settings_by_name("military_unit")

	def get_date_of_event(self):
		return self.get_report_settings_by_name("date_of_event")

	def get_report_settings_by_name(self, name):
		rep_settings = self.get_report_settings()
		if name in rep_settings:
			return rep_settings[name]
		else:
			return ""

	def get_commander_generic(self, settings_key, empty_placeholder, declension_type, is_short_name):
		c_name = f"[ФИО {empty_placeholder}]"
		c_rank = f"[ЗВАНИЕ {empty_placeholder}]"
		c_position = f"[ДОЛЖНОСТЬ {empty_placeholder}]"

		rep_settings = self.get_report_settings()
		found = settings_key in rep_settings

		if found:
			commander = rep_settings[settings_key]
			c_name = self.get_person_name_declension(commander["name"], declension_type)
			if is_short_name:
				c_name = self.get_person_name_short_format_1(c_name)
			c_rank = commander["rank"]
			c_guard = commander["is_guard"]
			c_rank_declension = get_word_declension(self.get_morph(), c_rank, declension_type)
			if c_guard:
				c_rank = "гвардии " + c_rank_declension
			else:
				c_rank = c_rank_declension
			c_position = glue_number_string(commander["position"].lower())
			if "заместитель" in c_position:
				c_position = self.word_replacement("заместитель", c_position, declension_type)
			else:
				if "командир" in c_position:
					c_position = self.word_replacement("командир", c_position, declension_type)
				else:
					c_position = get_word_declension(self.get_morph(), c_position, declension_type)
		return {"name": c_name, "rank": c_rank, "position": c_position, "found": found}

	# TODO private
	def word_replacement(self, word_to_replace, whole_string, declension_type):
		return whole_string.replace(word_to_replace,
		                            get_word_declension(self.get_morph(), word_to_replace, declension_type))

	def get_current_year(self):
		return datetime.date.today().year

	def get_service_started_str(self):
		service_started = self.get_report_settings()["service_started"]
		if service_started is None:
			return ""
		if isinstance(service_started, datetime.datetime):
			return f"{service_started.day}.{service_started.month}.{service_started.year}"
		return service_started

	def get_service_started_str_year(self):
		service_started = self.get_service_started_str()
		if "." in service_started:
			return service_started.split(".")[2]
		return "[УКАЖИТЕ ГОД НАЧАЛА СЛУЖБЫ]"

	def get_soldier_address(self, declension_type):
		s_info = self.get_soldier_info()
		sq = replace_with_glue(get_words_declension(self.get_morph(), s_info.squad, declension_type).strip())
		pl = replace_with_glue(get_words_declension(self.get_morph(), s_info.platoon, declension_type).strip())
		cm = replace_with_glue(get_words_declension(self.get_morph(), s_info.company, declension_type).strip())
		return f"{sq} {pl} {cm}"
