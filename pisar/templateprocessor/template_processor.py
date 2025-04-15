import os

from docx import Document

from classes.document_in_report import MODEL_OUTPUT_FOLDER, MODEL_CURRENT_SOLDIER
from helpers.log_helper import log
from templateprocessor.ru_date_event import RuDateEvent
from templateprocessor.ru_date_event_short import RuDateEventShort
from templateprocessor.ru_dob import RuDob
from templateprocessor.ru_dob_short import RuDobShort
from templateprocessor.ru_sold_additional_attributes import RuSoldAdditionalAttributes
from templateprocessor.ru_sold_citizenship import RuSoldCitizenship
from templateprocessor.ru_sold_education import RuSoldEducation
from templateprocessor.ru_sold_father import RuSoldFather
from templateprocessor.ru_sold_fio import RuSoldFio
from templateprocessor.ru_sold_fio_short1 import RuSoldFioShort1
from templateprocessor.ru_sold_home_address import RuSoldHomeAddress
from templateprocessor.ru_sold_home_address_real import RuSoldHomeAddressReal
from templateprocessor.ru_sold_marital_status import RuSoldMaritalStatus
from templateprocessor.ru_sold_mother import RuSoldMother
from templateprocessor.ru_sold_passport import RuSoldPassport
from templateprocessor.ru_sold_personal_perks import RuSoldPersonalPerks
from templateprocessor.ru_sold_phone import RuSoldPhone
from templateprocessor.ru_sold_place_of_birth import RuSoldPlaceOfBirth
from templateprocessor.ru_sold_position import RuSoldPosition
from templateprocessor.ru_sold_rank import RuSoldRank
from templateprocessor.ru_sold_registration_office import RuSoldRegistrationOffice
from templateprocessor.ru_sold_siblings import RuSoldSiblings
from templateprocessor.ru_sold_signs import RuSoldSigns
from templateprocessor.ru_sold_spouse import RuSoldSpouse
from templateprocessor.ru_sold_sr import RuSoldSr
from templateprocessor.ru_sold_tatoo import RuSoldTatoo
from templateprocessor.ru_sold_unique import RuSoldUnique

MODEL_BOX_FOLDER = "box_path"


class TemplateProcessor:
    def __init__(self, data_model, pers_storage):
        self.data_model = data_model
        self.pers_storage = pers_storage
        self.box_folder = self.data_model[MODEL_BOX_FOLDER]
        self.output_folder = os.path.join(self.data_model[MODEL_OUTPUT_FOLDER], "box")
        self.replacements = self.get_replacements()

    def process(self):
        if not os.path.exists(self.box_folder):
            os.makedirs(self.box_folder)
        if not os.path.exists(self.output_folder):
            os.makedirs(self.output_folder)
        log(f"Каталог с исходными документами: {self.box_folder}")
        log(f"Результаты здесь: {self.output_folder}")
        for file_name in os.listdir(self.box_folder):
            if not file_name.endswith(".docx"):
                continue
            full_path_input = os.path.join(self.box_folder, file_name)
            s_info = self.data_model[MODEL_CURRENT_SOLDIER]
            fn = file_name.replace(".docx", "") + f" ({s_info.full_name}).docx"
            full_path_output = os.path.join(self.output_folder, fn)
            self.process_document(full_path_input, full_path_output)

    def process_document(self, full_path_input, full_path_output):
        f = open(full_path_input, "rb")
        document = Document(f)
        f.close()

        self.process_document_runs(document)
        self.process_document_tables(document)

        document.save(full_path_output)
        log(f"Документ сохранен: {full_path_output}")

    def process_document_tables(self, document):
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    replacement_result = self.replace_all_placeholders(cell.text)
                    if replacement_result != cell.text:
                        cell.text = replacement_result

    def process_document_runs(self, document):
        for paragraph in document.paragraphs:
            if len(paragraph.runs) < 2:
                continue

            # merge runs, prepare them for replacement operations
            run_anchor = None
            for index in range(1, len(paragraph.runs)):
                # on start
                if run_anchor is None:
                    run_anchor = paragraph.runs[index - 1]
                run_this = paragraph.runs[index]
                if len(run_this.text) == 0:
                    continue

                if run_this.bold == run_anchor.bold and run_this.italic == run_anchor.italic and run_this.underline == run_anchor.underline:
                    # merge them
                    run_anchor.text = run_anchor.text + run_this.text
                    run_this.clear()
                else:
                    run_anchor = paragraph.runs[index]

            # replacement routines
            for run in paragraph.runs:
                ln = len(str(run.text))
                if ln > 0:
                    replacement_result = self.replace_all_placeholders(run.text)
                    if replacement_result != run.text:
                        run.text = replacement_result

    def replace_all_placeholders(self, run_text):
        current_text = ""
        has_changed = True
        while has_changed:
            for repl in self.replacements:
                if len(current_text) == 0:
                    current_text = run_text
                repl_text = repl.find(current_text)
                has_changed = len(repl_text) > 0
                if has_changed:
                    current_text = repl_text

        return current_text

    def get_replacements(self):
        result = [RuDob(self.data_model, self.pers_storage, "{СОЛД-ДР-ПОЛН}"),
                  RuSoldFio(self.data_model, self.pers_storage, "{СОЛД-ФИО;"),
                  RuSoldRank(self.data_model, self.pers_storage, "{СОЛД-ЗВАНИЕ;"),
                  RuSoldSr(self.data_model, self.pers_storage, "{СОЛД-ШР;"),
                  RuSoldPosition(self.data_model, self.pers_storage, "{СОЛД-ДОЛЖНОСТЬ;"),
                  RuDateEvent(self.data_model, self.pers_storage, "{ДАТА-СОБЫТИЯ}"),
                  RuDateEventShort(self.data_model, self.pers_storage, "{ДАТА-СОБЫТИЯ-КР}"),
                  RuSoldFioShort1(self.data_model, self.pers_storage, "{СОЛД-ИО-КР;"),
                  RuDobShort(self.data_model, self.pers_storage, "{СОЛД-ДР-КР}"),
                  RuSoldUnique(self.data_model, self.pers_storage, "{СОЛД-ЛН}"),
                  RuSoldRegistrationOffice(self.data_model, self.pers_storage, "{СОЛД-РВК}"),
                  RuSoldCitizenship(self.data_model, self.pers_storage, "{СОЛД-ГРАЖДАНСТВО}"),
                  RuSoldEducation(self.data_model, self.pers_storage, "{СОЛД-ОБРАЗОВАНИЕ}"),
                  RuSoldMaritalStatus(self.data_model, self.pers_storage, "{СОЛД-СЕМ-ПОЛОЖЕНИЕ}"),
                  RuSoldFather(self.data_model, self.pers_storage, "{СОЛД-ОТЕЦ}"),
                  RuSoldMother(self.data_model, self.pers_storage, "{СОЛД-МАТЬ}"),
                  RuSoldSpouse(self.data_model, self.pers_storage, "{СОЛД-ЖЕНА}"),
                  RuSoldSiblings(self.data_model, self.pers_storage, "{СОЛД-БРАТ-СЕСТРА}"),
                  RuSoldPassport(self.data_model, self.pers_storage, "{СОЛД-ПАСПОРТ}"),
                  RuSoldHomeAddress(self.data_model, self.pers_storage, "{СОЛД-АДРЕС-ПРОПИСКИ}"),
                  RuSoldHomeAddressReal(self.data_model, self.pers_storage, "{СОЛД-АДРЕС-ФАКТ}"),
                  RuSoldPhone(self.data_model, self.pers_storage, "{СОЛД-ТЕЛЕФОН}"),
                  RuSoldPlaceOfBirth(self.data_model, self.pers_storage, "{СОЛД-МЕСТО-РЖ}"),
                  RuSoldSigns(self.data_model, self.pers_storage, "{СОЛД-ПРИМЕТЫ}"),
                  RuSoldTatoo(self.data_model, self.pers_storage, "{СОЛД-ТАТУ}"),
                  RuSoldAdditionalAttributes(self.data_model, self.pers_storage, "{СОЛД-ПРИМЕТЫ-ДОП}"),
                  RuSoldPersonalPerks(self.data_model, self.pers_storage, "{СОЛД-ОТЛИЧ-ЗНАКИ}"),
                  ]
        return result
