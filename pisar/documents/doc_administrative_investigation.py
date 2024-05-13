# ДОКУМЕНТ
# АДМИНИСТРАТИВНОЕ РАССЛЕДОВАНИЕ по факту самовольного оставления части
from docx.shared import Mm

from classes.paragraph_settings import ParagraphSettings
from classes.pers_full_name_settings import PersFullNameSettings
from documents.investigation_prototype import InvestigationPrototype


class DocAdministrativeInvestigation(InvestigationPrototype):
	def get_name(self):
		return "Административное расследование по факту самовольного оставления части"

	def get_name_for_file(self):
		return f"Административное расследование_СОЧ ({self.get_soldier_info().full_name}).docx"

	def get_rank_name(self, declension_type):
		s_info = self.get_soldier_info()
		rank = self.get_person_rank(s_info.rank, declension_type)
		full_name = self.get_person_name_declension(self.get_person_name_short_format_1(s_info.full_name), declension_type)
		return f"{rank} {full_name}"

	def render(self, custom_margins=None):
		sold_str = self.get_rank_name(1)

		self.doc_title = "АДМИНИСТРАТИВНОЕ РАССЛЕДОВАНИЕ"
		self.doc_about = "по факту самовольного оставления части"

		self.inventory_caption = "административного расследования"
		rows_content = [f"Медицинская характеристика на {sold_str}", f"Копия служебной карточки на {sold_str}", f"Служебная характеристика на {sold_str}", "Справка по форме ГУК МО РФ от 04.12.2006 г.", "Копия ведомости закрепления оружия <ПОДРАЗДЕЛЕНИЕ>", "Копия именного списка для вечерних поверок <ПОДРАЗДЕЛЕНИЕ>", f"Копия ведомости об ознакомлении со статьями УК РФ {sold_str}", f"Выписки из приказов командира войсковой части {self.get_military_unit()} на {sold_str}", f"Копия паспорта {sold_str}", "Заключение административного разбирательства"]
		self.add_inventory_list(rows_content)

		self.report3_conclusion = "самовольно покинул расположение части, не уведомив вышестоящее командование"

		self.conclusion_p2_letter_to_parents = True
		self.conclusion_materials = "административного расследования"
		self.conclusion_fact = "самовольного оставления части"
		self.conclusion_action_performed = "административное расследование по факту самовольного оставления части"
		self.conclusion_p2_action = "самовольно покинул расположение части, не уведомив вышестоящее командование"

		self.conclusion_punishment = "самовольно покинул расположение части, за что в соответствии со статьей «337» Уголовного кодекса Российской Федерации предусматривается уголовная ответственность"

		settings = PersFullNameSettings(2, False, False, True, False, True, False)
		sold_str = self.get_person_full_str(settings)

		txt = f"3. Передать материалы административного расследования по факту самовольного оставления части {sold_str} в военную прокуратуру г.Донецка для дальнейшего принятия процессуального решения."
		self.conclusion_punishment_points.append(txt)

		settings = PersFullNameSettings(1, False, False, True, False, True, False)
		sold_str = self.get_person_full_str(settings)
		if len(sold_str) > 2:
			sold_str = sold_str[0].upper() + sold_str[1:]
		date_formatted = self.get_date_format_1(self.get_date_of_event())
		txt = f"4. {sold_str} снять с котлового довольствия с {date_formatted}."
		self.conclusion_punishment_points.append(txt)

		txt = f"5. {sold_str} снять с финансового обеспечения с {date_formatted}."
		self.conclusion_punishment_points.append(txt)

		super().render()

	def report1_page(self):
		self.add_paragraph(f"Командиру войсковой части {self.get_military_unit()}", self.align_right_settings)
		self.add_empty_paragraphs(2)
		self.add_paragraph("Рапорт", self.align_center_settings)
		self.add_empty_paragraphs(2)
		settings = PersFullNameSettings(2, False, False, True, True, False, False)
		sold_str = self.get_person_full_str(settings)
		date_str = self.get_date_format_1(self.get_date_of_event())
		self.add_paragraph(f"Настоящим докладываю, что {date_str} во время утренней поверки в <ПОДРАЗДЕЛЕНИЕ> 2 стрелкового батальона был выявлен факт самовольного оставления части {sold_str}.", self.ident_align_justify_settings)

		self.add_paragraph(
			f"Проведена поисковая работа, однако к должному результату данные мероприятия не привели.",
			self.ident_align_justify_settings)

		self.add_paragraph(
			f"Прошу Вашего указания на проведения административного расследования по данному факту.",
			self.ident_align_justify_settings)

		self.add_empty_paragraphs(3)
		self.officer_report_footer("commander_2_level", True)

	def report2_page(self):
		self.add_paragraph("Командиру 2 стрелкового батальона", self.align_right_settings)
		self.add_empty_paragraphs(2)
		self.add_paragraph("Рапорт", self.align_center_settings)
		self.add_empty_paragraphs(2)

		date_str = self.get_date_format_1(self.get_date_of_event())
		settings = PersFullNameSettings(2, False, False, True, True, False, False)
		sold_str = self.get_person_full_str(settings)
		settings = PersFullNameSettings(1, False, False, False, False, False, False, address_required=False)
		sold_str2 = self.get_person_full_str(settings)
		self.add_paragraph(f"Настоящим докладываю, что {date_str} в ходе установления причины самовольного оставления части {sold_str}, я в личном деле взял номер мобильного телефона матери {sold_str2} и позвонил ей с целью установления места пребывания {sold_str2}.", self.ident_align_justify_settings)

		self.add_paragraph(
			f"В ходе телефонного звонка я сообщил его матери, что он совершил уголовное преступление, за которое понесет уголовную ответственность согласно требованиям Уголовного кодекса Российской Федерации. После чего довел старшему поисковой команды <ДОЛЖНОСТЬ И ФИО> адрес возможного нахождения {sold_str2} для проведения поисковых мероприятий.",
			self.ident_align_justify_settings)

		self.add_empty_paragraphs(3)
		self.officer_report_footer("commander_1_level", True, "по")

	def report4_page(self):
		self.add_paragraph(f"Командиру <РОТА>", self.align_right_settings)
		self.add_paragraph(f"2 стрелкового батальона", self.align_right_settings)
		self.add_empty_paragraphs(2)
		self.add_paragraph("Рапорт", self.align_center_settings)
		self.add_empty_paragraphs(1)
		settings = PersFullNameSettings(0, False, False, True, False, False, False)
		sold_str = self.get_person_full_str(settings)
		self.add_paragraph(f"Настоящим докладываю, что {sold_str} сдал оружие <ТИП И НОМЕР> в комнату хранения оружия <РОТА> 2 стрелкового батальона, что подтверждается записями в книге выдачи оружия и боеприпасов <РОТА> 2 стрелкового батальона.", self.ident_align_justify_settings)

		self.add_empty_paragraphs(3)
		self.add_paragraph("<ВСТАВЬТЕ ДОЛЖНОСТНОЕ ЛИЦО ПО ШАБЛОНУ С ПРЕДЫДУЩЕЙ СТРАНИЦЫ>", self.align_left_settings)
		return True

	def conclusion_page(self):
		s_info = self.get_soldier_info()

		self.add_paragraph(f"Командиру войсковой части {self.get_military_unit()}", self.align_right_settings)
		self.add_empty_paragraphs(2)
		self.add_paragraph("ЗАКЛЮЧЕНИЕ", self.bold_center_settings)
		self.add_paragraph(f"по материалам {self.conclusion_materials}", self.bold_center_settings)
		self.add_paragraph(f"по факту {self.conclusion_fact}", self.bold_center_settings)
		self.add_paragraph(f"военнослужащим 2 стрелкового батальона войсковой части {self.get_military_unit()}",
		                   self.bold_center_settings)

		self.add_paragraph(f"{self.get_person_rank(s_info.rank, 2)} {self.get_person_name_instr(s_info.full_name)}",
		                   self.bold_center_settings)
		self.add_empty_paragraphs(1)

		self.add_paragraph_left_right(self.get_date_of_event(), "г. Донецк")

		paragraph_settings = ParagraphSettings()
		paragraph_settings.first_line_indent = Mm(12.5)
		paragraph_settings.align_justify = True
		paragraph_settings.line_spacing = 0.88

		# TODO ВРИО -> временно исполняющим обязанности. long_position?
		commander = self.get_commander_generic("commander_2_level", "КОМАНДИРА", 2, False)
		settings = PersFullNameSettings(2, False, False, True, True, True, False)
		sold_str = self.get_person_full_str(settings)
		self.add_paragraph(f"Мной, командиром 2 стрелкового батальона войсковой части {self.get_military_unit()} {commander['rank']} {commander['name']}, проведено административное расследование по факту самовольного оставления воинской части {sold_str}.", paragraph_settings)

		settings = PersFullNameSettings(0, False, False, True, False, True, False)
		sold_str = self.get_person_full_str(settings)
		date_str = self.get_date_format_1(self.get_date_of_event())
		self.add_paragraph(f"Так, {date_str} {sold_str} самовольно покинул расположение 2 стрелкового батальона не уведомив вышестоящее командование.", paragraph_settings)

		settings = PersFullNameSettings(0, False, False, False, False, False, False, address_required=False)
		sold_str = self.get_person_full_str(settings)
		self.add_paragraph(
			f"По результатам административного расследования установлено, что {sold_str}, без каких-либо уважительных причин, самовольно покинул расположение 2 стрелкового батальона, убыл в неизвестном направлении и отсутствует по настоящее время. Организованные поиски указанного военнослужащего к положительному результату не привели.",
			paragraph_settings)

		self.add_paragraph("Причинами данного правонарушения явились:", paragraph_settings)

		txt = "невыполнение требований статьи 144, 145 Устава Внутренней Службы Вооруженных Сил Российской Федерации в " \
		      "части, касающейся воспитания, поддержания воинской дисциплины, морально–психологического состояния " \
		      "личного состава в роте, а также в части касающееся знаний деловых и морально–психологических качеств и " \
		      "особенностей всех военнослужащих роты, постоянного проведения с ними индивидуальной работы по воинскому " \
		      "воспитанию"

		commander_company_text1 = self.get_commander_company_full_str(2, False)
		self.add_paragraph(f"{txt} {commander_company_text1};", paragraph_settings)

		# TODO use dynamic part
		commander_platoon = self.get_commander_platoon_full_str(2)
		txt = "невыполнение требований статьи 152, 153 Устава Внутренней Службы Вооруженных Сил Российской Федерации в " \
		      "части, касающейся воспитания, поддержания воинской дисциплины, морально–психологического состояния во " \
		      f"взводе {commander_platoon};"
		self.add_paragraph(txt, paragraph_settings)

		commander_squad = self.get_commander_generic_full_str("commander_squad", 2, False)
		txt = f"невыполнение требований статей 156, 157 Устава Внутренней Службы Вооружённых Сил Российской Федерации, в части касающейся воспитания, поддержания воинской дисциплины, укрепления морально-политического и психологического состояния подчиненного личного состава взводе {commander_squad};"
		self.add_paragraph(txt, paragraph_settings)

		settings = PersFullNameSettings(2, False, False, True, False, True, False)
		sold_str = self.get_person_full_str(settings)
		self.add_paragraph("невыполнение требований статьи 160, 161 Устава Внутренней Службы Вооруженных Сил "
		                   "Российской Федерации в части, касающейся точного и своевременного исполнения возложенных "
		                   f"на него обязанностей, поставленных задач {sold_str}, его личная недисциплинированность, низкие деловые, морально-политические и психологические качества военнослужащего.",
		                   paragraph_settings)

		settings = PersFullNameSettings(0, False, False, False, False, False, False, address_required=False)
		sold_str = self.get_person_full_str(settings)
		self.add_paragraph(f"Таким образом, {sold_str}, из-за личной недисциплинированности, без каких-либо уважительных причин, в результате ненадлежащего контроля за подчиненными и эффективной работы по их воспитанию со стороны должностных лиц 2 стрелкового батальона, {date_str} совершил самовольное оставление части в военное время либо в условиях вооруженного конфликта или ведения боевых действий, то есть действия по признакам преступления, предусмотренного в соответствии с требованием статьи 337 Уголовного кодекса Российской Федерации.", paragraph_settings)

		p1 = self.add_paragraph("В целях недопущения правонарушений, связанных с уклонением военнослужащих от исполнения обязанностей военной службы и самовольным оставлением части в дальнейшем, привлечением к ответственности виновных должностных лиц, ", paragraph_settings)
		runner = p1.add_run("ПРЕДЛАГАЮ:")
		runner.bold = True

		self.add_paragraph(f"1. Командирам подразделений 2 стрелкового батальона организовать доведения до всего личного состава сведений об уголовной ответственностиза преступления, связанные с уклонением от исполнения обязанностей военной службы в условиях вооруженного конфликта или ведения боевых действий, в соответствии с требованием Уголовного кодекса Российской Федерации.", paragraph_settings)

		txt1 = "За невыполнение требований  статей 144, 145 Устава внутренней службы Вооруженных Сил Российской Федерации, в части касающейся организации и проведения им работы по повседневному воспитанию, поддержанию воинской дисциплины, укреплению морально-политического и психологического состояния подчиненного личного состава"
		commander_company_text2 = self.get_commander_company_full_str(3, False)
		txt2 = "провести дополнительные занятия в роте по ознакомлению со статьями УК РФ и наказаниями за их нарушение."
		self.add_paragraph(f"2. {txt1} {commander_company_text2} {txt2}", paragraph_settings)

		# TODO dynamic
		txt1 = "За невыполнение требований статей 152, 153 Устава внутренней службы Вооруженных Сил Российской Федерации, в части касающейся воспитания, поддержания воинской дисциплины, укрепления морально-политического и психологического состояния солдат подчиненного взвода,"
		commander_platoon_text = self.get_commander_platoon_full_str(3)
		txt2 = ", строго указать на исполнение служебных и должностных обязанностей"
		self.add_paragraph(f"3. {txt1} {commander_platoon_text}{txt2}", paragraph_settings)

		commander_squad = self.get_commander_generic_full_str("commander_squad", 3, False)
		txt1 = "За невыполнение требований статей 156, 157 Устава Внутренней Службы Вооружённых Сил Российской Федерации, в части касающейся воспитания, поддержания воинской дисциплины, укрепления морально-политического и психологического состояния подчиненного личного состава взвода"
		txt2 = ", строго указать на низкую дисциплину в отделении."
		self.add_paragraph(f"4. {txt1} {commander_squad}{txt2}", paragraph_settings)

		settings = PersFullNameSettings(2, False, False, True, False, False, False)
		sold_str = self.get_person_full_str(settings)
		self.add_paragraph(
			f"5. Командиру 2 стрелкового батальона организовать передачу установленным порядком материалов административного расследования по факту самовольного оставления воинской части {sold_str} в военную прокуратуру города Донецка для дальнейшего принятия процессуального решения.",
			paragraph_settings)

		self.add_empty_paragraphs(3)

		self.officer_report_footer("commander_2_level")


