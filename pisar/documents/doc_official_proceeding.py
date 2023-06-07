# ДОКУМЕНТ
# СЛУЖЕБНОЕ РАЗБИРАТЕЛЬСТВО по факту грубого дисциплинарного проступка
from docx.shared import Mm

from classes.document_in_report import DocumentInReport
from classes.paragraph_settings import ParagraphSettings


class DocOfficialProceeding(DocumentInReport):
	def get_name(self):
		return "Служебное разбирательство по факту грубого дисциплинарного проступка"

	def get_name_for_file(self):
		return f"Служебное_Разбирательство_ГДП_СОЧ ({self.get_soldier_info().full_name}).docx"

	def render(self):
		# TODO order and list of documents should be customizable
		self.render_page_title()
		self.word_document.add_page_break()
		self.inventory_page()
		self.word_document.add_page_break()
		# TODO tabs in reports
		self.report1_page()
		self.word_document.add_page_break()
		self.report2_page()
		self.word_document.add_page_break()
		self.report3_page()
		self.word_document.add_page_break()
		self.conclusion_page()
		self.word_document.add_page_break()
		self.summary_page()

		super().render()

	# Титульный лист (стр 1)
	def render_page_title(self):
		self.add_empty_paragraphs(13)
		self.add_paragraph("СЛУЖЕБНОЕ РАЗБИРАТЕЛЬСТВО", self.bold_center_settings)
		self.add_empty_paragraphs(1)
		self.add_paragraph("по факту грубого дисциплинарного проступка", self.bold_center_settings)

		sold_str = self.get_person_full_str(2, True, True, False, False, True, False)

		self.add_paragraph(
			sold_str,
			self.bold_center_settings)
		self.add_empty_paragraphs(18)

		# TODO year should be changeable
		# TODO finish this piece
		# ds = rep_settings["date_started"]
		# df = rep_settings["date_finished"]

		# if ds is not None:

		self.add_paragraph("Начато «     » _________ 2023 г.", self.align_right_settings)
		self.add_paragraph("Окончено «     » _________ 2023 г.", self.align_right_settings)

	# Опись (стр 2)
	def inventory_page(self):
		rep_settings = self.get_report_settings()
		commander = rep_settings["commander_1_level"]

		self.add_paragraph("О П И С Ь", self.align_center_settings)
		self.add_paragraph("документов, находящихся в материалах служебного разбирательства",
		                   self.align_center_settings)
		self.add_empty_paragraphs(1)

		captions = ["№ п/п", "Название документа", "Номер листов"]
		rows_data = ["Рапорт ВРИО командира 2 СБ"
			, "Рапорт ВРИО ЗКБ по ВПР 2 СБ"
			, "Рапорт ВРИО командира 5 СР 2 СБ"
			, "Объяснение [звание и ФИО]"
			, "Объяснение [звание и ФИО]"
			, "Объяснение [звание и ФИО]"
			, "Протокол о грубом дисциплинарном проступке"
			, "Акт о невозможности получения копии протокола о грубом дисциплинарном проступке"
			,
			         "Акт о невозможности взять объяснение по факту действий совершенных военнослужащим, в которых усматривается преступление против военной службы"
			, "Служебная характеристика"
			, "Заключение служебного разбирательства"]

		self.add_table(12, captions, rows_data)
		self.add_empty_paragraphs(1)
		self.add_paragraph("Опись составил:", self.align_left_settings)
		self.add_paragraph(commander["position"], self.align_center_settings)
		self.add_paragraph(self.get_person_rank(commander["rank"], 0), self.align_center_settings)
		self.add_paragraph(self.get_person_name_short_format_1(commander["name"]), self.align_right_settings)

	# Рапорт 1 (стр 3)
	def report1_page(self):
		rep_settings = self.get_report_settings()
		commander = rep_settings["commander_2_level"]
		date_of_event = rep_settings["date_of_event"]

		self.add_paragraph("Командиру войсковой части " + rep_settings["military_unit"], self.align_right_settings)
		self.add_empty_paragraphs(2)
		self.add_paragraph("Рапорт", self.align_center_settings)
		self.add_empty_paragraphs(1)
		ps = "Настоящим докладываю, что"
		ps = f"{ps} {self.get_date_format_1(rep_settings['date_of_event'])} "

		sold_str = self.get_person_full_str(0, False, False, True, True, False, False)

		ps = ps + sold_str + ", отсутствовал на месте несения службы более 4 (четырех) часов, служебные обязанности не выполняет."
		self.add_paragraph(ps, self.ident_align_justify_settings)
		self.add_paragraph("Прошу Вашего указания на проведение служебного разбирательства по данному факту.",
		                   self.ident_align_justify_settings)
		self.add_empty_paragraphs(3)
		self.officer_report_footer(commander, False)

	# Рапорт 2 (стр 4)
	def report2_page(self):
		rep_settings = self.get_report_settings()
		commander = rep_settings["commander_1_level"]
		date_of_event = rep_settings["date_of_event"]

		self.add_paragraph("Командиру 2 стрелкового батальона", self.align_right_settings)
		self.add_empty_paragraphs(2)
		self.add_paragraph("Рапорт", self.align_center_settings)
		self.add_empty_paragraphs(1)

		ps = "Настоящим докладываю, что в ходе проведения розыскных мероприятий розыскной группой из числа " \
		     "военнослужащих 2 стрелкового батальона под моим руководством место нахождения"

		sold_str = self.get_person_full_str(1, False, False, True, False, False, False)

		ps = f"{ps} {sold_str} не установлено."
		self.add_paragraph(ps, self.ident_align_justify_settings)

		self.add_paragraph("Опрос сослуживцев результата не дал, на телефонные звонки не отвечает.",
		                   self.ident_align_justify_settings)
		self.add_empty_paragraphs(3)
		self.officer_report_footer(commander, False)

	# Рапорт 3 (стр 5)
	def report3_page(self):
		rep_settings = self.get_report_settings()
		date_of_event = rep_settings["date_of_event"]

		self.add_paragraph("Командиру 2 стрелкового батальона", self.align_right_settings)
		self.add_empty_paragraphs(2)
		self.add_paragraph("Рапорт", self.align_center_settings)
		self.add_empty_paragraphs(1)

		ps = f"Настоящим докладываю, что {self.get_date_format_1(date_of_event)}"

		sold_str = self.get_person_full_str(0, False, False, True, False, False, False)

		ps = f"{ps} {sold_str} отсутствует на службе без уважительных причин, не уведомив об этом вышестоящее командование."
		self.add_paragraph(ps, self.ident_align_justify_settings)
		self.add_empty_paragraphs(2)

		cc_info = self.get_commander_company()
		self.officer_report_footer(cc_info, True)

	# Заключение (стр 6 и 7)
	def conclusion_page(self):
		s_info = self.get_soldier_info()
		rep_settings = self.get_report_settings()
		commander = rep_settings["commander_2_level"]
		date_of_event = rep_settings["date_of_event"]

		self.add_paragraph(f"Командиру войсковой части {rep_settings['military_unit']}", self.align_right_settings)
		self.add_empty_paragraphs(3)
		self.add_paragraph("ЗАКЛЮЧЕНИЕ", self.bold_center_settings)
		self.add_paragraph("по материалам служебного разбирательства", self.bold_center_settings)
		self.add_paragraph("по факту грубого дисциплинарного проступка", self.bold_center_settings)
		self.add_paragraph(f"военнослужащим 2 стрелкового батальона войсковой части {rep_settings['military_unit']}",
		                   self.bold_center_settings)

		self.add_paragraph(f"{self.get_person_rank(s_info.rank, 2)} {self.get_person_name_instr(s_info.full_name)}",
		                   self.bold_center_settings)
		self.add_empty_paragraphs(1)

		self.add_paragraph_left_right(date_of_event, "г. Донецк")
		self.add_empty_paragraphs(1)

		paragraph_settings = ParagraphSettings()
		paragraph_settings.first_line_indent = Mm(12.5)
		paragraph_settings.align_justify = True
		paragraph_settings.line_spacing = 0.95

		# TODO ВРИО -> временно исполняющим обязанности. long_position?
		txt = f"Мной, временно исполняющим обязанности командира 2 стрелкового батальона войсковой части {rep_settings['military_unit']} "

		txt = txt + f"{self.get_person_rank(commander['rank'], 2)} {self.get_person_name_instr(commander['name'])},"
		txt = txt + " проведено служебное разбирательство по факту грубого дисциплинарного проступка"
		sold_str = self.get_person_full_str(2, False, False, True, True, True, False)
		self.add_paragraph(f"{txt} {sold_str}.", paragraph_settings)

		txt = f"В ходе проведения служебного разбирательства было установлено, что {self.get_date_format_1(rep_settings['date_of_event'])}"
		sold_str = self.get_person_full_str(0, False, False, True, False, False, False)
		txt = f"{txt} {sold_str} отсутствовал в месте несения военной службы без уважительных причин более 4 (четырех) часов подряд в течение установленного ежедневного служебного времени, тем самым совершил грубый дисциплинарный проступок."
		self.add_paragraph(txt, paragraph_settings)

		txt = "Проведенные розыскные мероприятия, опрос сослуживцев, поиск в лечебных учреждениях, " \
		       "военных комендатурах, а также отделениях полиции"
		txt = f"{txt} {self.get_person_rank(s_info.rank, 1)} {self.get_person_name_gent(s_info.full_name)} результата не дали, установить причины " \
		       f"отсутствия военнослужащего, а также его местонахождение не удалось."

		self.add_paragraph(txt, paragraph_settings)
		self.add_paragraph("Причинами данного правонарушения явились:", paragraph_settings)

		txt = "невыполнение требований статьи 144, 145 Устава Внутренней Службы Вооруженных Сил Российской Федерации в " \
		       "части, касающейся воспитания, поддержания воинской дисциплины, морально–психологического состояния " \
		       "личного состава в роте, а также в части касающееся знаний деловых и морально–психологических качеств и " \
		       "особенностей всех военнослужащих роты, постоянного проведения с ними индивидуальной работы по воинскому " \
		       "воспитанию"

		commander_company_text1 = self.get_commander_company_full_str(2)

		self.add_paragraph(f"{txt} {commander_company_text1};", paragraph_settings)

		# TODO use dynamic part
		commander_platoon = self.get_commander_platoon_full_str(2)
		txt = "невыполнение требований статьи 152, 153 Устава Внутренней Службы Вооруженных Сил Российской Федерации в " \
		      "части, касающейся воспитания, поддержания воинской дисциплины, морально–психологического состояния во " \
		      f"взводе командиром {commander_platoon};"

		self.add_paragraph(txt, paragraph_settings)

		sold_str = self.get_person_full_str(1, False, False, True, False, True, False)
		self.add_paragraph("невыполнение требований статьи 160, 161 Устава Внутренней Службы Вооруженных Сил "
		                   "Российской Федерации в части, касающейся точного и своевременного исполнения возложенных "
		                   f"на него обязанностей, поставленных задач и личная недисциплинированность {sold_str}.",
		                   paragraph_settings)

		self.add_paragraph(
			f"Исходя из материала служебного разбирательства следует, что {self.get_person_rank(s_info.rank, 0)} {s_info.full_name} "
			"отсутствовал в месте несения военной службы без уважительных причин более 4 (четырех) часов подряд в "
			"течение установленного ежедневного служебного времени, что в соответствии с Приложением 7 к "
			"Дисциплинарному уставу Вооруженных Сил Российской Федерации является грубым дисциплинарным проступком.",
			paragraph_settings)

		p1 = self.add_paragraph("На основании вышеизложенного ", self.ident_align_justify_settings)
		runner = p1.add_run("ПРЕДЛАГАЮ:")
		runner.bold = True

		commander_company_text2 = self.get_commander_company_full_str(3)

		txt = "за невыполнение требований статьи 144, 145 Устава Внутренней Службы Вооруженных Сил Российской " \
		       "Федерации в части, касающейся воспитания, поддержания воинской дисциплины, морально–психологического " \
		       "состояния личного состава в роте, а также в части касающееся знаний деловых и морально–психологических " \
		       "качеств и особенностей всех военнослужащих роты, постоянного проведения с ними индивидуальной работы " \
		       "по воинскому воспитанию, объявить ВЫГОВОР."
		self.add_paragraph(f"1. {commander_company_text2} {txt}", paragraph_settings)

		# TODO dynamic
		commander_platoon_text = self.get_commander_platoon_full_str(3)
		txt = f"2. Командиру {commander_platoon_text} за невыполнение требований статьи 152, 153 Устава Внутренней Службы Вооруженных Сил Российской Федерации в части, касающейся воспитания, поддержания воинской дисциплины, морально–психологического состояния во взводе, объявить СТРОГИЙ ВЫГОВОР."
		self.add_paragraph(txt, paragraph_settings)

		sold_str = self.get_person_full_str(1, False, False, True, False, True, False)
		txt = f"3. {sold_str} за грубый дисциплинарный проступок самовольное оставление места несения службы более 4 (четырех) часов, ПРЕДУПРЕДИТЬ О НЕПОЛНОМ СЛУЖЕБНОМ СООТВЕТСТВИИ."
		self.add_paragraph(txt, paragraph_settings)
		self.add_empty_paragraphs(2)

		self.officer_report_footer(commander, False)

	def summary_page(self):
		# TODO border around text
		self.add_empty_paragraphs(10)
		self.add_paragraph("В данном разбирательстве пронумеровано, прошито", self.align_center_settings)
		self.add_paragraph("и скреплено печатью ____ лист___.", self.align_center_settings)
		self.add_empty_paragraphs(1)
		self.officer_report_footer(self.get_report_settings()["commander_1_level"], False)

	# TODO find better way then is_name_shrunk
	def officer_report_footer(self, commander, is_shrunk):
		nm = commander["name"]
		rnk = commander["rank"]
		if not is_shrunk:
			nm = self.get_person_name_short_format_1(nm)
			rnk = self.get_person_rank(rnk, 0)

		rep_settings = self.get_report_settings()
		date_of_event = rep_settings["date_of_event"]
		self.add_paragraph(commander["position"], self.align_center_settings)
		self.add_paragraph(rnk, self.align_center_settings)
		self.add_paragraph_left_right(f"{date_of_event} г.", nm)
