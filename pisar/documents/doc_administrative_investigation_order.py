from docx.shared import Pt

from classes.document_in_report import DocumentInReport
from classes.paragraph_settings import ParagraphSettings
from classes.pers_full_name_settings import PersFullNameSettings


class DocAdministrativeInvestigationOrder(DocumentInReport):
	def get_name(self):
		return "Приказ"

	def get_name_for_file(self):
		return f"Приказ_СОЧ ({self.get_soldier_info().full_name}).docx"

	def render(self, custom_margins=None):
		self.copy_text()

		self.add_paragraph("Для служебного пользования", self.align_right_settings)
		self.add_paragraph("Экз. №____", self.align_right_settings)
		ps_title = ParagraphSettings()
		ps_title.is_bold = True
		ps_title.align_center = True
		ps_title.font_size = Pt(16)
		self.add_paragraph("П Р И К А З", ps_title)
		self.add_paragraph(f"КОМАНДИРА ВОЙСКОВОЙ ЧАСТИ {self.get_military_unit()}", self.bold_center_settings)
		self.add_paragraph(f"«___» _________ {self.get_current_year()} г.№___ ", self.align_center_settings)
		self.add_paragraph("г. Донецк", self.align_center_settings)
		self.add_empty_paragraphs(1)

		r_n = self.get_r_n(2)
		self.add_paragraph("По факту самовольного оставления части", self.bold_center_settings)
		self.add_paragraph(f"{r_n}", self.bold_center_settings)
		self.add_paragraph("____________________________________________________", self.align_center_settings)
		self.add_empty_paragraphs(1)

		self.add_paragraph(f"Несмотря на меры, принимаемые командованием войсковой части {self.get_military_unit()} по профилактике нарушения воинской дисциплины, направленные на укрепление правопорядка, сплочения воинских коллективов, создания в них здоровой морально–психологической обстановки, способствующих успешному выполнению поставленных задач, проведение профилактической работы по недопущению подобных случаев, среди военнослужащих продолжает иметь место случаи уклонения от исполнения обязанностей военной службы.", self.ident_align_justify_settings)
		settings = PersFullNameSettings(0, False, False, True, True, True, False)
		sold_str = self.get_person_full_str(settings)
		self.add_paragraph(f"В ходе проведения административного расследования было установлено, что {self.get_date_format_1(self.get_date_of_event())} {sold_str} самовольно покинул расположение части, не уведомив вышестоящее командование.", self.ident_align_justify_settings)

		r_n = self.get_r_n(1)
		self.add_paragraph(f"Проведенные розыскные мероприятия, опрос сослуживцев, поиск в лечебных учреждениях, отделениях полиции {r_n} результата не дали, установить причины отсутствия военнослужащего, а также его местонахождение не удалось. Подготовлено и направлено письмо по адресу проживания родителей {r_n} о его самовольном оставления части.", self.ident_align_justify_settings)

		self.add_paragraph("Причинами данного правонарушения явились:", self.ident_align_justify_settings)

		txt = "невыполнение требований статьи 144, 145 Устава Внутренней Службы Вооруженных Сил Российской Федерации в " \
		      "части, касающейся воспитания, поддержания воинской дисциплины, морально–психологического состояния " \
		      "личного состава в роте, а также в части касающееся знаний деловых и морально–психологических качеств и " \
		      "особенностей всех военнослужащих роты, постоянного проведения с ними индивидуальной работы по воинскому " \
		      "воспитанию"

		commander_company_text1 = self.get_commander_company_full_str(2)

		self.add_paragraph(f"{txt} {commander_company_text1};", self.ident_align_justify_settings)

		# TODO use dynamic part
		commander_platoon = self.get_commander_platoon_full_str(2)
		self.add_paragraph("невыполнение требований статьи 152, 153 Устава Внутренней Службы Вооруженных Сил Российской Федерации в " \
		      "части, касающейся воспитания, поддержания воинской дисциплины, морально–психологического состояния во " \
		      f"взводе командиром {commander_platoon};", self.ident_align_justify_settings)

		settings = PersFullNameSettings(1, False, False, True, False, True, False)
		sold_str = self.get_person_full_str(settings)
		self.add_paragraph("невыполнение требований статьи 160, 161 Устава Внутренней Службы Вооруженных Сил "
		                   "Российской Федерации в части, касающейся точного и своевременного исполнения возложенных "
		                   f"на него обязанностей, поставленных задач и личная недисциплинированность {sold_str}.",
		                   self.ident_align_justify_settings)

		s_info = self.get_soldier_info()
		self.add_paragraph(
			f"Исходя из материала служебного разбирательства следует, что {self.get_person_rank(s_info.rank, 0)} {s_info.full_name} самовольно покинул расположение части, за что в соответствии со статьей «337» Уголовного кодекса Российской Федерации предусматривается уголовная ответственность.",
			self.ident_align_justify_settings)

		p1 = self.add_paragraph("На основании вышеизложенного, ", self.ident_align_justify_settings)
		runner = p1.add_run("ПРИКАЗЫВАЮ:")
		runner.bold = True

		commander_company_text2 = self.get_commander_company_full_str(3)

		txt = "за невыполнение требований статьи 144, 145 Устава Внутренней Службы Вооруженных Сил Российской " \
		      "Федерации в части, касающейся воспитания, поддержания воинской дисциплины, морально–психологического " \
		      "состояния личного состава в роте, а также в части касающееся знаний деловых и морально–психологических " \
		      "качеств и особенностей всех военнослужащих роты, постоянного проведения с ними индивидуальной работы " \
		      "по воинскому воспитанию, объявить ВЫГОВОР."
		self.add_paragraph(f"1. {commander_company_text2} {txt}", self.ident_align_justify_settings)

		# TODO dynamic
		commander_platoon_text = self.get_commander_platoon_full_str(3)
		txt = f"2. Командиру {commander_platoon_text} за невыполнение требований статьи 152, 153 Устава Внутренней Службы Вооруженных Сил Российской Федерации в части, касающейся воспитания, поддержания воинской дисциплины, морально–психологического состояния во взводе, объявить СТРОГИЙ ВЫГОВОР."
		self.add_paragraph(txt, self.ident_align_justify_settings)

		settings = PersFullNameSettings(2, False, False, True, False, True, False)
		sold_str = self.get_person_full_str(settings)
		txt = f"3. Передать материалы административного расследования по факту самовольного оставления части {sold_str} в военную прокуратуру г.Донецка для дальнейшего принятия процессуального решения."
		self.add_paragraph(txt, self.ident_align_justify_settings)

		settings = PersFullNameSettings(1, False, False, True, False, True, False)
		sold_str = self.get_person_full_str(settings)
		if len(sold_str) > 2:
			sold_str = sold_str[0].upper() + sold_str[1:]
		date_formatted = self.get_date_format_1(self.get_date_of_event())
		txt = f"4. {sold_str} снять с котлового довольствия с {date_formatted}."
		self.add_paragraph(txt, self.ident_align_justify_settings)

		txt = f"5. {sold_str} снять с финансового обеспечения с {date_formatted}."
		self.add_paragraph(txt, self.ident_align_justify_settings)

		self.add_paragraph(f"6. Контроль за исполнением настоящего приказа возложить на начальника штаба войсковой части {self.get_military_unit()}.", self.ident_align_justify_settings)
		self.add_paragraph(f"7. Приказ довести до личного состава в части касающейся.", self.ident_align_justify_settings)

		self.add_empty_paragraphs(1)
		self.officer_report_footer("commander_4_level")
		self.add_empty_paragraphs(2)
		self.officer_report_footer("commander_3_level")

		self.copy_correct_text()
		super().render()

	def copy_text(self):
		pass

	def copy_correct_text(self):
		pass

	def get_r_n(self, declension_type):
		s_info = self.get_soldier_info()
		rank = self.get_person_rank(s_info.rank, declension_type)
		name = self.get_person_name_declension(s_info.full_name, declension_type)
		return f"{rank} {name}"

	def officer_report_footer(self, key):
		commander = self.get_commander_generic(key, "КОМАНДИРА", 0, True)
		self.add_paragraph(commander["position"].upper(), self.bold_center_settings)
		self.add_paragraph(commander["rank"], self.bold_center_settings)
		self.add_paragraph(commander["name"], self.bold_right_settings)

