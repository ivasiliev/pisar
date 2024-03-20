from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from classes.custom_margins import CustomMargins
from classes.document_in_report import DocumentInReport
from classes.paragraph_settings import ParagraphSettings


class ApprovalSheetPrototype(DocumentInReport):
	def __init__(self, data_model):
		super().__init__(data_model)
		self.title = ""

	def get_name(self):
		return "Лист согласования"

	def get_name_for_file(self):
		return f"Лист согласования ({self.get_soldier_info().full_name}).docx"

	def render(self, custom_margins=None):
		s_info = self.get_soldier_info()
		rep_settings = self.get_report_settings()

		# TODO this code is SIMILAR in Performance_Characteristics

		line_spacing = 0.96
		par_set_center = ParagraphSettings()
		par_set_center.is_bold = True
		par_set_center.align_center = True
		par_set_center.line_spacing = line_spacing

		par_set_right = ParagraphSettings()
		par_set_right.align_right = True
		par_set_right.is_bold = True
		par_set_right.line_spacing = line_spacing

		self.add_paragraph("ЛИСТ СОГЛАСОВАНИЯ", self.bold_center_settings)
		self.add_paragraph(
			f"приказа командира войсковой части {self.get_military_unit()}",
			self.align_center_settings)
		rank = self.get_person_rank(s_info.rank, 2)
		name = self.get_person_name_declension(s_info.full_name, 2)
		self.add_paragraph("«" + self.title, self.align_center_settings)
		self.add_paragraph(f"{rank} {name}»", self.align_center_settings)
		self.add_empty_paragraphs(1)

		# approval table
		f_caption = ParagraphSettings()
		f_caption.font_size = Pt(12)
		table_settings = {"ps": f_caption, "cols_width": [115, 50]}
		captions = ["Воинское звание, инициалы, фамилия, и занимаемая должность лица, с которым согласовывался приказ",
		            "Подпись должностного лица и дата согласования"]
		rows = self.add_empty_rows(4)
		self.add_table(captions, rows, table_settings)
		self.add_empty_paragraphs(1)

		self.add_commander(rep_settings["commander_3_level"], self.get_military_unit(), par_set_center, par_set_right)

		self.add_empty_paragraphs(3)

		self.add_paragraph("ЛИСТ ДОВЕДЕНИЯ", self.bold_center_settings)
		self.add_paragraph(f"приказа командира войсковой части {self.get_military_unit()}", self.align_center_settings)

		# informed table
		captions = ["Воинское звание, инициалы, фамилия, и занимаемая должность лица, которым доводился приказ",
		            "Подпись должностного лица и дата ознакомления с приказом"]
		rows = self.add_empty_rows(8)
		self.add_table(captions, rows, table_settings)
		self.add_empty_paragraphs(1)

		self.add_commander(rep_settings["commander_3_level"], self.get_military_unit(), par_set_center, par_set_right)
		self.add_empty_paragraphs(5)

		ps = ParagraphSettings()
		ps.align_left = True
		ps.font_size = Pt(12)

		cmd = self.get_commander_generic("commander_1_level", "ИСПОЛНИТЕЛЯ", 0, True)
		self.add_paragraph(f"Исп. {cmd['rank']} {cmd['name']}", ps)
		self.add_paragraph("   т. АТС ID 856-23", ps)

		super().render(CustomMargins(20, 20, 10, 30))

	def add_empty_rows(self, how_many):
		arr = []
		for i in range(0, how_many):
			arr.append(["", ""])
		return arr

	def add_commander(self, commander_info, military_unit, par_set_center, par_set_right):
		c_name = self.get_person_name_short_format_1(commander_info["name"])
		c_rank = commander_info["rank"]
		if self.get_report_settings()["is_guard"]:
			c_rank = "гвардии " + c_rank
		c_position = commander_info["position"] + " " + military_unit
		self.add_paragraph(c_position.upper(), self.bold_center_settings)
		self.add_paragraph(c_rank, par_set_center)
		self.add_paragraph(c_name, par_set_right)
