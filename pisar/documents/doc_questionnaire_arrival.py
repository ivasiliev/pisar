from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from classes.document_in_report import DocumentInReport
from classes.paragraph_settings import ParagraphSettings
from helpers.text_helper import not_empty


class DocQuestArrival(DocumentInReport):
	def get_name(self):
		return "Анкета прибывшего в зону СВО"

	def get_name_for_file(self):
		return f"{self.get_name()} ({self.get_soldier_info().full_name}).docx"

	def render(self):
		right_10 = ParagraphSettings()
		right_10.font_size = Pt(10)
		right_10.align_right = True

		self.add_paragraph("Для служебного пользования", right_10)
		self.add_paragraph("Экз. № ___", right_10)

		table_captions = ["", "", ""]
		table_settings = {"ps": None, "cols_width": [30, 150, 150], "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT,
		                  "font_size": Pt(10)}

		s_info = self.get_soldier_info()
		rep_settings = self.get_report_settings()

		# relatives info
		relatives_info = self.add_relatives_item("father", "Отец", False)
		relatives_info += f"{self.add_relatives_item('mother', 'Мать', True)}"
		relatives_info += f"{self.add_relatives_item('siblings', 'Братья/Сёстры', True)}"
		relatives_info += f"{self.add_relatives_item('spouse', 'Жена', True)}"

		# prepare passport info
		passport_info = ""
		if not_empty(rep_settings["pass_rf"]):
			passport_info = f"паспорт РФ {rep_settings['pass_rf']} {rep_settings['pass_rf_issued']}"
		if not_empty(rep_settings["pass_dnr"]):
			passport_info = f"{passport_info}\nпаспорт ДНР {rep_settings['pass_dnr']} {rep_settings['pass_dnr_issued']}"
		if not_empty(rep_settings["pass_ukr"]):
			passport_info = f"{passport_info}\nпаспорт Украины {rep_settings['pass_ukr']}"
		if not_empty(rep_settings["pass_foreign"]):
			passport_info = f"{passport_info}\nзагранпаспорт {rep_settings['pass_foreign']}"

		row_data = [
			["", "Анкета на прибывшего в зону СВО\n", ""]
			, ["1.", "Фамилия Имя Отчество", s_info.full_name]
			, ["2.", "Дата рождения", s_info.get_dob()]
			, ["3.", "Место рождения", rep_settings["place_of_birth"]]
			, ["4.", "Условное наименование воинской части (места службы)", self.get_military_unit()]
			, ["5.", "Воинское звание", rep_settings["rank"]]
			, ["6.", "Воинская должность", rep_settings["position"]]
			, ["7.", "Документы удостоверяющие личность (серия, номер, дата выдачи и кем выдан)", passport_info]
			, ["8.", "Личный номер", s_info.unique]
			, ["9.", "Место жительства", rep_settings["home_address"]]
			, ["10.", "Контактный номер телефона", rep_settings["phone"]]
			, ["11.", "Сведения о близких родственниках:(степень родства, ФИО, дата рождения, место жительства, номер телефона)", relatives_info]
			, ["12.", "Командиры (начальники):(воинское звание, ФИО, место жительства, телефон и другая информация о должностных лицах для взаимодействия)", "< вставьте сведения>"]
			, ["13.", "Особенности внешности военнослужащего:(цвет волос (стрижка (наличие залысин), наличие бороды и бакенбардов), строение тела, описание лица (форма носа, цвет глаз, уши, подбородок), отличительные черты описать подробно). ФОТО обязательно (анфас, профиль и отдельно особенностей)", rep_settings["signs"]]
			, ["14.", "Личные приметы:(наличие с описанием татуировок, шрамов, (где, когда и при каких обстоятельствах получен), родимых пятен, ампутаций, протезов, описание зубов и т.д.). Тату, шрамы и родимые пятна – ФОТО (обязательно)", rep_settings["tatoo"]]
			, ["15.", "Индивидуальные отличительные признаки:(одежды, снаряжения, вооружения (размер одежды и обуви, какое имеет снаряжение и оружие, их номера, наличие особенностей и меток) ФОТО Дополнительно: информация о наличии нашивок, надписей, позывных на элементах снаряжения. ФОТО", rep_settings["personal_perks"]]
			, ["16.", "Другая значимая информация:(наличие и описание отдельных элементов на военнослужащем: наручные часы, очки (контактные линзы), обереги, амулеты, кольца (в том числе обручальные), перстни, цепочки (с указанием типа металла), шнурки, молитвенная атрибутика (крестики, иконки, карманные молитвенники). Наличие вредных привычек.Указать группу крови (при наличии информации).ФОТО отдельных элементов.", rep_settings["additional_attributes"]]
		]
		self.add_table(table_captions, row_data, table_settings)
		self.add_empty_paragraphs(1)

		left_10 = ParagraphSettings()
		left_10.font_size = Pt(10)
		left_10.align_left = True

		center_8 = ParagraphSettings()
		center_8.font_size = Pt(8)
		center_8.align_center = True

		self.add_paragraph("Анкету оформил: __________________________________________________________________", left_10)
		self.add_paragraph("__________________________________________________________________________________", left_10)
		self.add_paragraph("(должность, в/звание, подпись, ФИО должностного лица)", center_8)
		self.add_paragraph("«___»____________2024 г.", left_10)

		self.word_document.add_page_break()

		super().render()

	def add_relatives_item(self, key, title, is_new_str):
		result = ""
		rep_settings = self.get_report_settings()
		if not_empty(rep_settings[f"{key}_name"]):
			result = f"{title}: {rep_settings[f'{key}_name']}\n{rep_settings[f'{key}_address']}\n{rep_settings[f'{key}_phone']}"
			if is_new_str:
				result = f"\n{result}"
		return result
