from docx.shared import Pt

from classes.document_in_report import DocumentInReport
from classes.paragraph_settings import ParagraphSettings


class DocApprovalSheet(DocumentInReport):
	def get_name(self):
		return "Лист согласования"

	def get_name_for_file(self):
		return f"Лист согласования_СОЧ ({self.get_soldier_info().full_name}).docx"

	def render(self):
		s_info = self.get_soldier_info()

		self.add_paragraph("ЛИСТ СОГЛАСОВАНИЯ", self.bold_center_settings)
		self.add_paragraph(
			f"проекта приказа командира войсковой части {self.get_military_unit()}",
			self.align_center_settings)
		rank = self.get_person_rank(s_info.rank, 2)
		name = self.get_person_name_declension(s_info.full_name, 2)
		self.add_paragraph("«По факту самовольного оставления части", self.align_center_settings)
		self.add_paragraph(f"{rank} {name}»", self.align_center_settings)
		self.add_empty_paragraphs(1)

		self.add_approval_table()
		self.add_empty_paragraphs(1)

		self.add_paragraph("ЛИСТ ДОВЕДЕНИЯ", self.bold_center_settings)
		self.add_paragraph(f"приказа командира войсковой части {self.get_military_unit()}", self.align_center_settings)
		self.add_approval_table()
		self.add_empty_paragraphs(1)

		ps = ParagraphSettings()
		ps.align_left = True
		ps.font_size = Pt(12)

		cmd = self.get_commander_generic("commander_1_level", "ИСПОЛНИТЕЛЯ")
		self.add_paragraph(f"Исп. {cmd['rank']} {cmd['name']}", ps)
		self.add_paragraph("   т. АТС ID 85623", ps)

		super().render()

	def add_empty_rows(self, how_many):
		arr = []
		for i in range(0, how_many):
			arr.append(["\n", ""])
		return arr

	def add_approval_table(self):
		captions = ["Воинское звание, инициалы, фамилия и должность лица, с которым согласовывался проект приказа",
		            "Замечания по проекту. Подпись должностного лица и дата согласования"]
		rows = self.add_empty_rows(6)
		self.add_table(captions, rows)
