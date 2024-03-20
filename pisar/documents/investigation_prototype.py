from docx.shared import Mm

from classes.document_in_report import DocumentInReport
from classes.paragraph_settings import ParagraphSettings
from classes.pers_full_name_settings import PersFullNameSettings


class InvestigationPrototype(DocumentInReport):
	def __init__(self, data_model):
		super().__init__(data_model)

		self.doc_title = ""
		self.doc_about = ""

		self.inventory_caption = ""
		self.inventory_list = []
		# TODO should we use here positions from config files?
		rows_content = ["Рапорт командира 2 СБ", "Рапорт ЗКБ по ВПР 2 СБ",
		                       "Рапорт командира 5 СР"]
		for i in range(0, 3):
			rows_content.append("Объяснение [звание и ФИО]")
		self.add_inventory_list(rows_content)

		self.report1_action = ""
		self.report1_request = ""

		self.report3_conclusion = ""

		self.conclusion_materials = ""
		self.conclusion_fact = ""
		self.conclusion_action_performed = ""
		self.conclusion_p2_action = ""
		self.conclusion_punishment = ""
		self.conclusion_p2_letter_to_parents = False
		# 1 and 2 are constants; add here points from 3 and further
		self.conclusion_punishment_points = []

	def render(self, custom_margins=None):
		self.render_page_title()
		self.word_document.add_page_break()
		self.inventory_page()
		self.word_document.add_page_break()
		self.report1_page()
		self.word_document.add_page_break()
		self.report2_page()
		self.word_document.add_page_break()
		self.report3_page()
		self.word_document.add_page_break()
		self.conclusion_page()
		self.word_document.add_page_break()
		# self.summary_page()

		super().render()

	# Титульный лист (стр 1)
	def render_page_title(self):
		self.add_empty_paragraphs(13)
		self.add_paragraph(self.doc_title, self.bold_center_settings)
		self.add_empty_paragraphs(1)
		self.add_paragraph(self.doc_about, self.bold_center_settings)

		settings = PersFullNameSettings(2, True, True, False, False, True, False)
		sold_str = self.get_person_full_str(settings)

		self.add_paragraph(
			sold_str,
			self.bold_center_settings)
		self.add_empty_paragraphs(18)

		current_year = self.get_current_year()

		self.add_paragraph(f"Начато «     » _________ {current_year} г.", self.align_right_settings)
		self.add_paragraph(f"Окончено «     » _________ {current_year} г.", self.align_right_settings)

	# Опись (стр 2)
	def inventory_page(self):
		self.add_paragraph("О П И С Ь", self.align_center_settings)
		self.add_paragraph(f"документов, находящихся в материалах {self.inventory_caption}",
		                   self.align_center_settings)
		self.add_empty_paragraphs(1)

		captions = ["№ п/п", "Название документа", "Номер листов"]

		capt_settings = {"ps": None, "cols_width": [10, 120, 35]}
		self.add_table(captions, self.inventory_list, capt_settings)
		self.add_empty_paragraphs(1)
		self.add_paragraph("Опись составил:", self.align_left_settings)
		self.add_empty_paragraphs(1)
		self.officer_report_footer("commander_1_level")

	# Рапорт 1 (стр 3)
	def report1_page(self):
		self.add_paragraph(f"Командиру войсковой части {self.get_military_unit()}", self.align_right_settings)
		self.add_empty_paragraphs(2)
		self.add_paragraph("Рапорт", self.align_center_settings)
		self.add_empty_paragraphs(1)
		ps = "Настоящим докладываю, что"
		ps = f"{ps} {self.get_date_format_1(self.get_date_of_event())} "

		settings = PersFullNameSettings(0, False, False, True, True, False, False)
		sold_str = self.get_person_full_str(settings)

		ps = ps + sold_str + f", {self.report1_action}."
		self.add_paragraph(ps, self.ident_align_justify_settings)
		if len(self.report1_request) > 0:
			self.add_paragraph(f"Прошу Вашего указания на {self.report1_request}.",
			                   self.ident_align_justify_settings)
		self.add_empty_paragraphs(3)
		self.officer_report_footer("commander_2_level")

	# Рапорт 2 (стр 4)
	def report2_page(self):
		self.add_paragraph("Командиру 2 стрелкового батальона", self.align_right_settings)
		self.add_empty_paragraphs(2)
		self.add_paragraph("Рапорт", self.align_center_settings)
		self.add_empty_paragraphs(1)

		ps = "Настоящим докладываю, что в ходе проведения розыскных мероприятий розыскной группой из числа " \
		     "военнослужащих 2 стрелкового батальона под моим руководством место нахождения"

		settings = PersFullNameSettings(1, False, False, True, False, False, False)
		sold_str = self.get_person_full_str(settings)

		ps = f"{ps} {sold_str} не установлено."
		self.add_paragraph(ps, self.ident_align_justify_settings)

		self.add_paragraph("Опрос сослуживцев результата не дал, на телефонные звонки не отвечает.",
		                   self.ident_align_justify_settings)
		self.add_empty_paragraphs(3)
		self.officer_report_footer("commander_1_level")

	# Рапорт 3 (стр 5)
	def report3_page(self):
		self.add_paragraph("Командиру 2 стрелкового батальона", self.align_right_settings)
		self.add_empty_paragraphs(2)
		self.add_paragraph("Рапорт", self.align_center_settings)
		self.add_empty_paragraphs(1)

		ps = f"Настоящим докладываю, что {self.get_date_format_1(self.get_date_of_event())}"
		settings = PersFullNameSettings(0, False, False, True, False, False, False)
		sold_str = self.get_person_full_str(settings)
		self.add_paragraph(f"{ps} {sold_str} {self.report3_conclusion}.", self.ident_align_justify_settings)
		self.add_empty_paragraphs(2)

		self.officer_report_footer("commander_company")

	# Заключение (стр 6 и 7)
	def conclusion_page(self):
		s_info = self.get_soldier_info()

		self.add_paragraph(f"Командиру войсковой части {self.get_military_unit()}", self.align_right_settings)
		self.add_empty_paragraphs(2)
		self.add_paragraph("ЗАКЛЮЧЕНИЕ", self.bold_center_settings)
		self.add_paragraph(f"по материалам {self.conclusion_materials}", self.bold_center_settings)
		self.add_paragraph(f"по факту {self.conclusion_fact}", self.bold_center_settings)
		self.add_paragraph(f"военнослужащим 2 стрелкового батальона войсковой части {self.get_military_unit()}",
		                   self.bold_center_settings)

		self.add_paragraph(f"{self.get_person_rank(s_info.rank, 2)} {self.get_person_name_instr(s_info.full_name)}",
		                   self.bold_center_settings)
		self.add_empty_paragraphs(1)

		self.add_paragraph_left_right(self.get_date_of_event(), "г. Донецк")

		paragraph_settings = ParagraphSettings()
		paragraph_settings.first_line_indent = Mm(12.5)
		paragraph_settings.align_justify = True
		paragraph_settings.line_spacing = 0.88

		# TODO ВРИО -> временно исполняющим обязанности. long_position?
		txt = f"Мной, командиром 2 стрелкового батальона войсковой части {self.get_military_unit()} "

		commander = self.get_commander_generic("commander_2_level", "КОМАНДИРА", 2, False)
		txt = txt + f"{commander['rank']} {commander['name']},"
		txt = txt + f" проведено {self.conclusion_action_performed}"
		settings = PersFullNameSettings(2, False, False, True, True, True, False)
		sold_str = self.get_person_full_str(settings)
		self.add_paragraph(f"{txt} {sold_str}.", paragraph_settings)

		txt = f"В ходе проведения {self.conclusion_materials} было установлено, что {self.get_date_format_1(self.get_date_of_event())}"
		settings = PersFullNameSettings(0, False, False, True, False, False, False)
		sold_str = self.get_person_full_str(settings)
		txt = f"{txt} {sold_str} {self.conclusion_p2_action}."
		self.add_paragraph(txt, paragraph_settings)

		txt = "Проведенные розыскные мероприятия, опрос сослуживцев, поиск в лечебных учреждениях, " \
		      "военных комендатурах, а также отделениях полиции"
		txt = f"{txt} {self.get_person_rank(s_info.rank, 1)} {self.get_person_name_gent(s_info.full_name)} результата не дали, установить причины " \
		      f"отсутствия военнослужащего, а также его местонахождение не удалось."
		if self.conclusion_p2_letter_to_parents:
			rank = self.get_person_rank(s_info.rank, 1)
			name = self.get_person_name_declension(s_info.full_name, 1)
			txt = f"{txt} Подготовлено и направлено письмо по адресу проживания родителей {rank} {name} о его самовольном оставлении части."

		self.add_paragraph(txt, paragraph_settings)
		self.add_paragraph("Причинами данного правонарушения явились:", paragraph_settings)

		txt = "невыполнение требований статьи 144, 145 Устава Внутренней Службы Вооруженных Сил Российской Федерации в " \
		      "части, касающейся воспитания, поддержания воинской дисциплины, морально–психологического состояния " \
		      "личного состава в роте, а также в части касающееся знаний деловых и морально–психологических качеств и " \
		      "особенностей всех военнослужащих роты, постоянного проведения с ними индивидуальной работы по воинскому " \
		      "воспитанию"

		commander_company_text1 = self.get_commander_company_full_str(2, False)
		self.add_paragraph(f"{txt} {commander_company_text1};", paragraph_settings)

		# TODO use dynamic part
		commander_platoon = self.get_commander_platoon_full_str(2)
		txt = "невыполнение требований статьи 152, 153 Устава Внутренней Службы Вооруженных Сил Российской Федерации в " \
		      "части, касающейся воспитания, поддержания воинской дисциплины, морально–психологического состояния во " \
		      f"взводе {commander_platoon};"
		self.add_paragraph(txt, paragraph_settings)

		commander_squad = self.get_commander_generic_full_str("commander_squad", 2, False)
		txt = f"невыполнение требований статей 156, 157 Устава Внутренней Службы Вооружённых Сил Российской Федерации, в части касающейся воспитания, поддержания воинской дисциплины, укрепления морально-политического и психологического состояния подчиненного личного состава взводе {commander_squad};"
		self.add_paragraph(txt, paragraph_settings)

		settings = PersFullNameSettings(1, False, False, True, False, True, False)
		sold_str = self.get_person_full_str(settings)
		self.add_paragraph("невыполнение требований статьи 160, 161 Устава Внутренней Службы Вооруженных Сил "
		                   "Российской Федерации в части, касающейся точного и своевременного исполнения возложенных "
		                   f"на него обязанностей, поставленных задач и личная недисциплинированность {sold_str}.",
		                   paragraph_settings)

		self.add_paragraph(
			f"Исходя из материала служебного разбирательства следует, что {self.get_person_rank(s_info.rank, 0)} {s_info.full_name} {self.conclusion_punishment}.",
			paragraph_settings)

		p1 = self.add_paragraph("На основании вышеизложенного ", self.ident_align_justify_settings)
		runner = p1.add_run("ПРЕДЛАГАЮ:")
		runner.bold = True

		txt1 = "За невыполнение требований  статей 144, 145 Устава внутренней службы Вооруженных Сил Российской Федерации, в части касающейся организации и проведения им работы по повседневному воспитанию, поддержанию воинской дисциплины, укреплению морально-политического и психологического состояния подчиненного личного состава"
		commander_company_text2 = self.get_commander_company_full_str(3, False)
		txt2 = "провести дополнительные занятия в роте по ознакомлению со статьями УК РФ и наказаниями за их нарушение."
		self.add_paragraph(f"1. {txt1} {commander_company_text2} {txt2}", paragraph_settings)

		# TODO dynamic
		txt1 = "За невыполнение требований статей 152, 153 Устава внутренней службы Вооруженных Сил Российской Федерации, в части касающейся воспитания, поддержания воинской дисциплины, укрепления морально-политического и психологического состояния солдат подчиненного взвода,"
		commander_platoon_text = self.get_commander_platoon_full_str(3)
		txt2 = ", строго указать на исполнение служебных и должностных обязанностей"
		self.add_paragraph(f"2. {txt1} {commander_platoon_text}{txt2}", paragraph_settings)

		commander_squad = self.get_commander_generic_full_str("commander_squad", 3, False)
		txt1 = "За невыполнение требований статей 156, 157 Устава Внутренней Службы Вооружённых Сил Российской Федерации, в части касающейся воспитания, поддержания воинской дисциплины, укрепления морально-политического и психологического состояния подчиненного личного состава взвода"
		txt2 = ", строго указать на низкую дисциплину в отделении."
		self.add_paragraph(f"3. {txt1} {commander_squad}{txt2}", paragraph_settings)

		for prg in self.conclusion_punishment_points:
			self.add_paragraph(prg, paragraph_settings)

		self.add_empty_paragraphs(1)

		self.officer_report_footer("commander_2_level")

	def summary_page(self):
		# TODO border around text
		self.add_empty_paragraphs(10)
		self.add_paragraph("В данном разбирательстве пронумеровано, прошито", self.align_center_settings)
		self.add_paragraph("и скреплено печатью ____ лист___.", self.align_center_settings)
		self.add_empty_paragraphs(1)
		self.officer_report_footer("commander_1_level")

	def officer_report_footer(self, key, need_capitalize=False):
		commander = self.get_commander_generic(key, "КОМАНДИРА", 0, True)
		pos = commander["position"]
		if need_capitalize:
			pos = commander["position"].capitalize()
		self.add_paragraph(pos, self.align_center_settings)
		self.add_paragraph(commander["rank"], self.align_center_settings)
		self.add_paragraph_left_right(f"{self.get_date_of_event()} г.", commander["name"])

	def add_inventory_list(self, rows_content):
		num_row = len(self.inventory_list) + 1
		for r in rows_content:
			self.inventory_list.append([str(num_row), r, ""])
			num_row = num_row + 1
