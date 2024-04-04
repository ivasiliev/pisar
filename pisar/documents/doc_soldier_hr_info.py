from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from classes.document_in_report import DocumentInReport
from classes.pers_full_name_settings import PersFullNameSettings


class DocSoldierHrCard(DocumentInReport):
	def get_name(self):
		return "Справка"

	def get_name_for_file(self):
		return f"Справка ({self.get_soldier_info().full_name}).docx"

	def render(self, custom_margins=None):
		rep_settings = self.get_report_settings()
		s_info = self.get_soldier_info()

		self.add_paragraph("С П Р А В К А", self.bold_center_settings)
		# TODO same code goes to helper
		tokens = s_info.full_name.split(" ")
		name = ""
		if len(tokens) != 3:
			name = "[ФИО в правильном формате]"
		else:
			for t_ind in range(0, len(tokens)):
				s = tokens[t_ind]
				if t_ind == 0:
					s = s.upper()
				name = name + " " + s
		self.add_paragraph(name.strip(), self.bold_center_settings)

		settings = PersFullNameSettings(0, False, False, True, False, True, False, False, False)
		sold_str = self.get_person_full_str(settings)
		self.add_paragraph(sold_str, self.align_center_settings)
		self.add_empty_paragraphs(1)

		self.add_left_right_with_captions("Дата рождения", "Место рождения", s_info.get_dob(), rep_settings["place_of_birth"])
		self.add_left_right_with_captions("Образование", "Окончил (когда, что)", rep_settings["education"], rep_settings["graduation_place"])
		self.add_left_right_with_captions("Специальность", "", rep_settings["specialization"], "")
		self.add_left_right_with_captions("Какими иностранными языками владеет", "Является ли депутатом", rep_settings["foreign_languages"], rep_settings["government_authority"])
		self.add_left_right_with_captions("Имеет ли государственные награды (какие)", "Был ли за границей (когда, где)", rep_settings["awards"], rep_settings["foreign_countries_visited"])

		self.add_bold_caption("Воинское звание", self.get_person_rank(s_info.rank, 0))
		self.add_bold_caption("Контракт", f"Мобилизован ДНР с {self.get_service_started_str()}")

		marital_status = rep_settings["marital_status"]
		father_name = rep_settings["father_name"]
		mother_name = rep_settings["mother_name"]
		family_status = ""
		if len(marital_status) > 0:
			family_status = marital_status
		if father_name is not None and len(father_name) > 0:
			family_status = family_status + f"\nОтец – {father_name}"
		if mother_name is not None and len(mother_name) > 0:
			family_status = family_status + f"\nМать – {mother_name}"
		family_status = family_status.rstrip()
		self.add_bold_caption("Семейное положение", family_status)
		self.add_empty_paragraphs(1)

		self.add_paragraph("РАБОТА В ПРОШЛОМ", self.bold_center_settings)
		self.add_paragraph(rep_settings["occupation"], self.align_center_settings)
		self.add_empty_paragraphs(2)

		hr_officer = self.get_commander_generic("hr_officer", "СПЕЦИАЛИСТА ОТДЕЛА КАДРОВ", 0, True)
		self.add_paragraph(hr_officer["position"], self.align_center_settings)
		self.add_paragraph(f"войсковой части {self.get_military_unit()}", self.align_center_settings)
		self.add_paragraph(hr_officer["rank"], self.align_center_settings)
		self.add_paragraph(hr_officer["name"], self.align_right_settings)

		super().render()

	def add_left_right_with_captions(self, left_caption, right_caption, left_content, right_content):
		table = self.word_document.add_table(rows=1, cols=2)
		table.alignment = WD_TABLE_ALIGNMENT.CENTER
		cells_captions = table.rows[0].cells
		# captions
		p_left_caption = cells_captions[0].paragraphs[0]  # cells_captions[0].add_paragraph()
		p_left_caption.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
		runner = p_left_caption.add_run(left_caption)
		runner.bold = True
		p_left_caption.add_run(f"\n{left_content}")
		p_right_caption = cells_captions[1].paragraphs[0]  # cells_captions[1].add_paragraph()
		p_right_caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
		runner = p_right_caption.add_run(right_caption)
		runner.bold = True
		p_right_caption.add_run(f"\n{right_content}")

	def add_bold_caption(self, caption, text):
		table = self.word_document.add_table(rows=1, cols=2)
		table.alignment = WD_TABLE_ALIGNMENT.CENTER
		cells = table.rows[0].cells
		p = cells[0].paragraphs[0]
		p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
		runner = p.add_run(caption)
		runner.bold = True

		p = cells[1].paragraphs[0]
		p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
		p.add_run(text)

		self.set_column_width(table, 0, 50)
		self.set_column_width(table, 1, 100)





