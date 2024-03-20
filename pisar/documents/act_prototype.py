from docx.shared import Mm

from classes.document_in_report import DocumentInReport
from classes.paragraph_settings import ParagraphSettings

ACT_TITLE = "act_title"
ACT_TEXT = "act_text"


# base class for Acts. Don't use it directly.
class ActPrototype(DocumentInReport):
	def render(self, custom_margins=None):
		rep_settings = self.get_report_settings()

		paragraph_settings = ParagraphSettings()
		paragraph_settings.left_indent = Mm(115)
		self.add_paragraph("«УТВЕРЖДАЮ»", paragraph_settings)

		paragraph_settings = ParagraphSettings()
		paragraph_settings.left_indent = Mm(95)
		self.add_empty_paragraphs(1)
		self.add_paragraph(f"Командир 2 стрелкового батальона", paragraph_settings)
		self.add_paragraph(f"войсковой части {rep_settings['military_unit']}", paragraph_settings)

		comm_2 = self.get_commander_generic("commander_2_level", "КОМАНДИРА", 0, True)
		self.add_paragraph(comm_2["rank"], paragraph_settings)

		paragraph_settings = ParagraphSettings()
		paragraph_settings.right_indent = Mm(31)
		paragraph_settings.align_left = True
		self.add_paragraph(comm_2["name"], self.align_right_settings)
		self.add_empty_paragraphs(1)
		self.add_paragraph("«___»____________2024 г.", self.align_right_settings)
		self.add_empty_paragraphs(3)
		self.add_paragraph("АКТ", self.bold_center_settings)
		self.add_paragraph(self.data_model[ACT_TITLE], self.bold_center_settings)
		self.add_empty_paragraphs(1)

		self.add_paragraph(self.data_model[ACT_TEXT], self.ident_align_justify_settings)
		self.add_empty_paragraphs(2)
		# self.add_paragraph("«___»____________2024 г.", self.align_left_settings)
		# self.add_empty_paragraphs(2)

		self.print_commander("commander_company_deputy_1")
		self.add_empty_paragraphs(2)

		self.print_commander("commander_company")
		self.add_empty_paragraphs(2)

		self.print_commander("commander_1_level", "по военно")

		super().render()

	def print_commander(self, key, breaker: str = ""):
		comm = self.get_commander_generic(key, "КОМАНДИРА", 0, True)
		pos = comm["position"].capitalize()
		if len(breaker) > 0 and breaker in pos:
			ind = pos.index(breaker)
			pos1 = pos[0:ind].strip()
			pos2 = pos[ind:].strip()
			self.add_paragraph(pos1, self.align_center_settings)
			self.add_paragraph(pos2, self.align_center_settings)
		else:
			self.add_paragraph(pos, self.align_center_settings)
		self.add_paragraph(comm["rank"], self.align_center_settings)
		self.add_paragraph(comm["name"], self.align_right_settings)
