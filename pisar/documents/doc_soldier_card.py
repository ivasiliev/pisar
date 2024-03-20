from docx.shared import Pt

from classes.document_in_report import DocumentInReport
from classes.paragraph_settings import ParagraphSettings


class DocSoldierCard(DocumentInReport):
	def get_name(self):
		return "Служебная карточка"

	def get_name_for_file(self):
		return f"Служебная_карточка ({self.get_soldier_info().full_name}).docx"

	def render(self, custom_margins=None):
		s_info = self.get_soldier_info()

		ps_28 = ParagraphSettings()
		ps_28.font_size = Pt(28)
		ps_28.align_center = True
		ps_28.is_bold = True

		ps_26 = ParagraphSettings()
		ps_26.font_size = Pt(26)
		ps_26.align_center = True
		ps_26.is_bold = True

		ps_20 = ParagraphSettings()
		ps_20.font_size = Pt(20)
		ps_20.align_center = True

		ps_16 = ParagraphSettings()
		ps_16.font_size = Pt(16)
		ps_16.align_center = True
		ps_16.is_italic = True

		rank = self.get_person_rank(s_info.rank, 0)
		tokens = s_info.full_name.split(" ")
		if len(tokens) != 3:
			print(f"Неверный формат ФИО военнослужащего '{s_info.full_name}'. Создание документа прервано.")
			return
		surname = tokens[0]
		name = tokens[1]
		second_name = tokens[2]

		# title page
		self.add_empty_paragraphs(10)
		self.add_paragraph("СЛУЖЕБНАЯ КАРТОЧКА", ps_28)
		self.add_empty_paragraphs(4)
		self.add_paragraph(rank, ps_16)
		self.add_empty_paragraphs(4)
		self.add_paragraph(surname, ps_26)
		self.add_empty_paragraphs(4)
		self.add_paragraph(f"{name} {second_name}", ps_20)
		self.word_document.add_page_break()

		# page 1
		ps_28_unbold = ParagraphSettings()
		ps_28_unbold.font_size = Pt(28)
		ps_28_unbold.align_center = True

		ps_18 = ParagraphSettings()
		ps_18.font_size = Pt(18)
		ps_18.align_center = True

		self.add_paragraph("СЛУЖЕБНАЯ КАРТОЧКА", ps_28_unbold)
		self.add_empty_paragraphs(1)
		self.add_paragraph("___________________________________________________________________", self.align_center_settings)
		self.print_special_paragraph(f"1. Воинская должность__", s_info.position)
		self.add_paragraph("___________________________________________________________________", self.align_center_settings)
		self.print_special_paragraph(f"2. Звание__", rank)
		self.print_special_paragraph(f"3. Фамилия, имя, отчество__", s_info.full_name)
		yss = self.get_service_started_str_year()
		self.print_special_paragraph(f"4. С какого года на военной службе__", yss)
		self.add_empty_paragraphs(1)
		self.add_paragraph("ПООЩРЕНИЯ", ps_18)
		captions = ["За что", "Вид поощрения", "Когда применено\n(дата, № приказа)", "Кем поощрен"]
		capt_ps = ParagraphSettings()
		capt_ps.is_bold = True
		capt_settings = {"ps": capt_ps, "cols_width": None}
		rows = self.add_empty_rows(len(captions), 25)  # 26
		self.add_table(captions, rows, capt_settings)
		self.word_document.add_page_break()

		# page 2
		self.add_paragraph("ДИСЦИПЛИНАРНЫЕ ВЗЫСКАНИЯ", ps_18)
		captions = ["Основание\nприменения\nвзыскания", "Когда\nсовершен\nпроступок", "Вид\nвзыскания", "Когда\nприменено\n(дата, № приказа)", "Кем\nприменено", "Когда приведено\nв исполнение", "Когда снято\n(кем или\nпо истечению срока)"]
		rows = self.add_empty_rows(len(captions), 35)  # 36
		capt_ps = ParagraphSettings()
		capt_ps.is_bold = True
		capt_ps.font_size = Pt(12)
		capt_settings = {"ps": capt_ps, "cols_width": None}
		self.add_table(captions, rows, capt_settings)
		self.word_document.add_page_break()

		# page 3
		self.add_paragraph("ЛИСТ ОЗНАКОМЛЕНИЯ", ps_18)
		captions = ["Дата", "Фамилия и инициалы военнослужащего", "Подпись военнослужащего"]
		rows = self.add_empty_rows(len(captions), 35)  # 36
		capt_ps = ParagraphSettings()
		capt_ps.is_bold = True
		capt_settings = {"ps": capt_ps, "cols_width": None}
		self.add_table(captions, rows, capt_settings)

		super().render()

	def print_special_paragraph(self, left_part, right_part):
		p = self.add_paragraph(left_part, self.align_left_settings)
		runner = p.add_run(right_part)
		runner.underline = True
		p.add_run("____________________")

	def add_empty_rows(self, cols_count, rows_count):
		rows = []
		for i in range(0, rows_count):
			row = []
			for j in range(0, cols_count):
				row.append("")
			rows.append(row)
		return rows
