from docx.shared import Pt

from classes.document_in_report import DocumentInReport
from classes.paragraph_settings import ParagraphSettings


class DocOrientation(DocumentInReport):
	def get_name(self):
		return "Ориентировка"

	def get_name_for_file(self):
		return f"Ориентировка_СОЧ ({self.get_soldier_info().full_name}).docx"

	def render(self):
		ps_title = ParagraphSettings()
		ps_title.font_size = Pt(36)
		ps_title.align_center = True
		ps_title.is_bold = True

		self.add_paragraph("ОН ПРЕДАЛ РОДИНУ", ps_title)
		self.add_empty_paragraphs(1)

		self.add_paragraph("[ВСТАВЬТЕ ФОТОГРАФИЮ]", self.align_center_settings)

		ps_text = ParagraphSettings()
		ps_text.font_size = Pt(22)
		ps_text.align_center = True
		ps_text.is_bold = True

		s_info = self.get_soldier_info()
		tokens = s_info.full_name.split(" ")
		name = ""
		if len(tokens) != 3:
			name = "[ФИО в правильном формате]"
		else:
			for t_ind in range(0, len(tokens)):
				s = tokens[t_ind]
				if t_ind == 0:
					s = s.upper()
				name = name + " " + s

		self.add_empty_paragraphs(1)
		self.add_paragraph(name.strip(), ps_text)
		self.add_empty_paragraphs(1)
		self.add_paragraph(f"Самовольно оставил подразделение, бросил боевых товарищей, тем самым поставил под угрозу выполнение боевой задачи. Если увидите – сообщите о его местонахождении по номеру: {self.get_report_settings()['phone_contact_to_report']}", ps_text)

		super().render()

