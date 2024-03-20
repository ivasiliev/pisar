# ДОКУМЕНТ
# СЛУЖЕБНАЯ ХАРАКТЕРИСТИКА
from docx.shared import Mm
from docx.shared import Pt

from classes.document_in_report import DocumentInReport
from classes.paragraph_settings import ParagraphSettings
from classes.pers_full_name_settings import PersFullNameSettings
from helpers.text_helper import get_word_gent


class DocPerformanceCharacteristics(DocumentInReport):
	def get_name(self):
		return "Служебная характеристика"

	def get_name_for_file(self):
		return f"{self.get_name()} ({self.get_soldier_info().full_name}).docx"

	def render(self, custom_margins=None):
		s_info = self.get_soldier_info()
		rep_settings = self.get_report_settings()
		line_spacing = 0.96

		paragraph_settings = ParagraphSettings()
		paragraph_settings.font_size = Pt(14)
		paragraph_settings.is_bold = True
		paragraph_settings.align_center = True

		p_title = self.add_paragraph("СЛУЖЕБНАЯ ХАРАКТЕРИСТИКА", paragraph_settings)
		p_title.paragraph_format.space_after = Pt(10)

		paragraph_settings = ParagraphSettings()
		paragraph_settings.left_indent = Mm(65)
		paragraph_settings.align_justify = True
		paragraph_settings.line_spacing = line_spacing
		settings = PersFullNameSettings(1, False, False, True, True, False, False)
		sold_str = self.get_person_full_str(settings)
		nationality = get_word_gent(self.get_morph(), rep_settings["nationality"])
		education = rep_settings["education"]
		yss = self.get_service_started_str_year()
		txt = f"на {sold_str}, {nationality}, образование {education}, в ВС РФ с {yss} года."
		self.add_paragraph(txt, paragraph_settings)

		self.add_empty_paragraphs_spacing(1, line_spacing)  # 2

		position_str = get_word_gent(self.get_morph(), s_info.position)
		settings = PersFullNameSettings(0, False, False, False, False, True, False)
		sold_str = self.get_person_full_str(settings).strip()
		txt = f"За время прохождения службы в должности {position_str} {sold_str}, проявил себя неоднозначно, как военнослужащий, требующий контроля со стороны командования при исполнении поставленных задач."
		paragraph_settings = self.ident_align_justify_settings
		paragraph_settings.line_spacing = line_spacing
		self.add_paragraph(txt, self.ident_align_justify_settings)

		txt = "По предметам боевой подготовки показал удовлетворительные знания. При несении службы показал себя как " \
		      "не самый дисциплинированный военнослужащий. Ранее был несколько раз уличён в нежелании грамотно, " \
		      "четко, а также своевременно выполнять поставленные задачи. Однако стоит отметить, что подобное " \
		      "происходило не на постоянной основе, хоть и является закономерностью."
		self.add_paragraph(txt, paragraph_settings)

		txt = "Данный военнослужащий, пользуется средним авторитетом среди сослуживцев и командования, служебные " \
		      "отношения средние. С начальством не пререкается, однако на замечания реагирует не всегда, что при этом " \
		      "приводит к неправильным выводам."
		self.add_paragraph(txt, paragraph_settings)

		txt = "На практике полученные знания применить не старается. На здоровье жалоб не было, физически развит " \
		      "удовлетворительно, имеет высокую работоспособность. Трудности военной службы переносит терпимо. В " \
		      "полном объеме владеет профессиональными знаниями, но овладеть новыми знаниями не стремится."
		self.add_paragraph(txt, paragraph_settings)

		self.add_empty_paragraphs_spacing(1, line_spacing)

		par_set_center = ParagraphSettings()
		par_set_center.is_bold = True
		par_set_center.align_center = True
		par_set_center.line_spacing = line_spacing

		par_set_right = ParagraphSettings()
		par_set_right.align_right = True
		par_set_right.is_bold = True
		par_set_right.line_spacing = line_spacing

		self.add_commander(self.get_commander_company(), par_set_center, par_set_right)
		self.add_empty_paragraphs_spacing(1, line_spacing)

		self.add_commander(rep_settings["commander_2_level"], par_set_center, par_set_right)

		self.add_paragraph("Подпись командира батальона заверяю:", self.align_justify_settings)
		# self.add_empty_paragraphs_spacing(1, line_spacing)

		# TODO check if correct
		#comm3 = rep_settings["commander_3_level"]
		#comm3_pos = comm3["position"] + " " + self.get_military_unit()
		self.add_commander(rep_settings["commander_3_level"], par_set_center, par_set_right)

		self.add_paragraph("«___» _________ 2024 г.", self.bold_justify_settings)

		super().render()

	def add_commander(self, commander_info, par_set_center, par_set_right):
		c_name = self.get_person_name_short_format_1(commander_info["name"])
		c_rank = commander_info["rank"]
		if self.get_report_settings()["is_guard"] and not "гвардии" in c_rank:
			c_rank = "гвардии " + c_rank
		c_position = f"{commander_info['position']} войсковой части {self.get_military_unit()}"
		self.add_paragraph(c_position.upper(), self.bold_center_settings)
		self.add_paragraph(c_rank, par_set_center)
		self.add_paragraph(c_name, par_set_right)
