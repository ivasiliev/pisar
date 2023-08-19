from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from classes.document_in_report import DocumentInReport
from classes.paragraph_settings import ParagraphSettings
from classes.personnel_storage import EXCEL_DOCUMENT_SR
from helpers.text_helper import get_date_str_format2


class DocDiaryBirthdays(DocumentInReport):
	def get_name(self):
		return "Дни рождения"

	def get_name_for_file(self):
		return f"{self.get_name()}.docx"

	def render(self):
		f_12_left = ParagraphSettings()
		f_12_left.align_left = True
		f_12_left.font_size = Pt(12)

		f_caption = ParagraphSettings()
		f_caption.font_size = Pt(8)

		pers_storage = self.get_pers_storage()
		all_persons = pers_storage.get_all_persons(EXCEL_DOCUMENT_SR)

		captions = ["Фамилия и инициалы военнослужащего", "январь", "февраль", "март", "апрель", "май", "июнь", "июль", "август", "сентябрь", "октябрь", "ноябрь", "декабрь"]
		table_settings = {"cols_width": [100], "ps": f_caption, "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT,
		                  "font_size": Pt(8)}
		for r in range(1, 13):
			table_settings["cols_width"].append(40)
		row_data = []
		for pers in all_persons:
			if pers.dob is None:
				continue
			pers_row = [self.get_person_name_short_format_2(pers.full_name, 0)]
			m = pers.dob.month
			for r in range(1, 13):
				if m == r:
					pers_row.append(get_date_str_format2(pers.dob))
				else:
					pers_row.append("")
			row_data.append(pers_row)

		self.add_paragraph("6. Дополнительные сведения, необходимые для индивидуальной работы с военнослужащими", self.bold_center_settings)
		self.add_paragraph("1) Дни рождения военнослужащих", f_12_left)
		self.add_table(captions, row_data, table_settings)
		self.word_document.add_page_break()

		super().render()