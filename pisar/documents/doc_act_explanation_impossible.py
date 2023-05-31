from docx.shared import Mm

from classes.document_in_report import DocumentInReport
from classes.paragraph_settings import ParagraphSettings


class DocActExplanationImpossible(DocumentInReport):
	def get_name(self):
		return "Акт о невозможности взять объяснения по факту ГДП"

	def get_name_for_file(self):
		return f"{self.get_name()} ({self.get_soldier_info().full_name}).docx"

	def render(self):
		s_info = self.get_soldier_info()
		rep_settings = self.get_report_settings()

		paragraph_settings = ParagraphSettings()
		paragraph_settings.left_indent = Mm(115)
		paragraph_settings.is_bold = True
		self.add_paragraph("УТВЕРЖДАЮ", paragraph_settings)
		self.add_paragraph(f"Командир войсковой части {rep_settings['military_unit']}", self.bold_right_settings)

		paragraph_settings = ParagraphSettings()
		paragraph_settings.right_indent = Mm(31)
		paragraph_settings.align_right = True
		paragraph_settings.is_bold = True
		comm_3 = rep_settings["commander_3_level"]
		self.add_paragraph(self.get_person_rank(comm_3["rank"], 0), paragraph_settings)
		self.add_paragraph(self.get_person_name_short_format_1(comm_3["name"]), self.bold_right_settings)
		self.add_empty_paragraphs(1)
		self.add_paragraph("«___»____________2023 г.", self.bold_right_settings)
		self.add_empty_paragraphs(3)
		self.add_paragraph("АКТ", self.bold_center_settings)
		self.add_paragraph("о невозможности взять объяснения по факту действий совершенных военнослужащим, в которых "
		                   "усматривается преступления против военной службы.", self.bold_center_settings)
		self.add_empty_paragraphs(1)

		sold_str = self.get_person_full_str(0, False, False, True, False)
		txt = f"Мы, нижеподписавшиеся, составили настоящий акт о том, что {sold_str} не может дать объяснения и не подписал составленный протокол о грубом дисциплинарном проступке, а также не получил копию по причине отсутствия."
		self.add_paragraph(txt, self.ident_align_justify_settings)
		self.add_empty_paragraphs(2)
		self.add_paragraph("«___»____________2023 г.", self.align_left_settings)
		self.add_empty_paragraphs(2)

		self.print_commander(rep_settings["commander_2_level"])
		self.add_empty_paragraphs(2)
		self.print_commander(rep_settings["commander_1_level"])
		self.add_empty_paragraphs(1)

		comm_company = self.get_commander_company()
		self.print_commander_routines(comm_company["name"], comm_company["rank"], comm_company["position"])

		super().render()

	def print_commander(self, comm):
		self.print_commander_routines(self.get_person_name_short_format_1(comm["name"]), self.get_person_rank(comm["rank"], 0), comm["position"])

	def print_commander_routines(self, c_name, c_rank, c_position):
		self.add_paragraph(c_position, self.align_center_settings)
		self.add_paragraph(c_rank, self.align_center_settings)
		self.add_paragraph(c_name, self.align_right_settings)
