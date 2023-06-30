from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Mm

from classes.document_in_report import DocumentInReport
from classes.paragraph_settings import ParagraphSettings
from classes.pers_full_name_settings import PersFullNameSettings


class DocOfficialProceedingProtocol(DocumentInReport):
	def get_name(self):
		return "Протокол по факту грубого дисциплинарного проступка"

	def get_name_for_file(self):
		return f"Протокол ГДП ({self.get_soldier_info().full_name}).docx"

	def render(self):
		sold = self.get_soldier_info()

		small_underline = ParagraphSettings()
		small_underline.underline = True
		small_underline.font_size = Pt(10)
		small_underline.align_center = True

		self.add_paragraph("ПРОТОКОЛ", self.bold_center_settings)
		self.add_paragraph("О ГРУБОМ ДИСЦИПЛИНАРНОМ ПРОСТУПКЕ", self.bold_center_settings)
		self.add_empty_paragraphs(1)
		runs = self.add_paragraph_left_right("«     » _________ 2023 г.", "населенный пункт г.Донецк")
		runs[1].underline = True

		commander1_full_str = self.get_commander_generic_full_str("commander_1_level", 0, "[ВСТАВЬТЕ СВЕДЕНИЯ О КОМАНДИРЕ]")
		self.add_paragraph_with_underline(commander1_full_str)
		self.add_paragraph("воинская должность, звание, ФИО лица составившего протокол", small_underline)

		self.add_paragraph_with_underline(f"Войсковая часть {self.get_military_unit()}")
		self.add_paragraph("сведения о военнослужащем: условное наименование воинской части (организации)", small_underline)

		settings = PersFullNameSettings(0, False, False, True, False, True, False, False)
		soldier_full_str = self.get_person_full_str(settings)
		self.add_paragraph_with_underline(soldier_full_str)
		self.add_paragraph("воинская должность, звание", small_underline)

		self.add_paragraph_with_underline(sold.full_name)
		self.add_paragraph("Фамилия, Имя, Отчество", small_underline)

		rep_settings = self.get_report_settings()
		birth_date = self.get_date_format_1(sold.get_dob())
		birth_place = rep_settings["place_of_birth"]
		home_address = rep_settings["home_address"]
		education = rep_settings["education"]
		nationality = rep_settings["nationality"]
		passport = rep_settings["passport"]
		marital_status = rep_settings["marital_status"]
		criminal_status = rep_settings["criminal_status"]

		self.add_paragraph_with_underline(f"{birth_date}, родился в {birth_place}")
		self.add_paragraph("год и место рождения", small_underline)

		self.add_paragraph_with_underline(f"проживает в {home_address}, {marital_status}")
		self.add_paragraph("место жительства (регистрации), семейное положение", small_underline)

		self.add_paragraph_with_underline(passport)
		self.add_paragraph("данные о документе, удостоверяющем личность (серия, номер, когда и кем выдан)", small_underline)

		self.add_paragraph_with_underline(f"образование {education}, {marital_status}, {nationality}")
		self.add_paragraph("иные сведения о военнослужащем", small_underline)

		self.add_paragraph_with_underline(criminal_status)
		self.add_paragraph("в том числе: привлекался ли ранее к дисциплинарной ответственности, когда и кем", small_underline)

		commander_company = self.get_commander_company_full_str(0)
		self.render_combined_row("Очевидцы: ", commander_company)
		self.add_paragraph("должности, места военной службы, звания ФИО лиц, которым известны обстоятельства, имеющие значение для правильного решения вопроса о привлечении военнослужащего", small_underline)

		self.add_paragraph("Обстоятельства совершения грубого дисциплинарного проступка:", self.align_justify_settings)
		date_event = self.get_date_of_event()
		settings = PersFullNameSettings(0, False, False, True, False, True, False)
		soldier_full_str = self.get_person_full_str(settings)
		self.add_paragraph_with_underline(f"{date_event} {soldier_full_str}")
		self.add_paragraph("воинская должность, звание", small_underline)

		self.add_paragraph_with_underline("отсутствовал в месте военной службы без уважительных причин более 4 (четырех) часов подряд в течение установленного ежедневного служебного времени")
		self.add_paragraph("дата, время, место и другие обстоятельства совершения грубого дисциплинарного проступка", small_underline)

		self.add_paragraph("Доказательства, подтверждающие наличия события грубого дисциплинарного проступка и виновности военнослужащего:", self.align_justify_settings)
		self.render_combined_row("Рапорт ", f"{commander_company}, объяснения сослуживцев, другие материалы служебного разбирательства")
		self.add_paragraph("перечисление доказательств: объяснения военнослужащего, привлекаемого к дисциплинарной ответственности, объяснения очевидцев, заключение и пояснение специалиста, документы, вещественные доказательства и др.", small_underline)

		settings = PersFullNameSettings(3, False, False, False, False, True, False, False, False)
		soldier_full_str = self.get_person_full_str(settings)
		self.render_combined_row("Военнослужащему ", soldier_full_str)
		self.add_paragraph("воинская должность", small_underline)

		settings = PersFullNameSettings(0, False, False, False, False, False, False, True, True, False)
		soldier_full_str = self.get_person_full_str(settings)
		self.add_paragraph_with_underline(soldier_full_str)
		self.add_paragraph("звание, ФИО", small_underline)

		pers_name_formatted = self.get_person_name_short_format_2(sold.full_name, 1)
		self.add_paragraph("разъяснены права и обязанности, предусмотренные законодательством Российской Федерации и общевоинскими уставами.", self.align_justify_settings)
		self.render_combined_row("Подпись военнослужащего: ", f"нет, в связи с отсутствием {pers_name_formatted}")
		self.render_combined_row("Объяснения военнослужащего, совершившего грубый дисциплинарный проступок: ", f"нет, в связи с отсутствием {pers_name_formatted}")
		self.render_combined_row("Смягчающие или отягчающие обстоятельства: ", "нет")
		self.render_combined_row("Причины и условия, способствовавшие совершению грубого дисциплинарного проступка: ", f"низкие морально-деловые качества {pers_name_formatted}")
		self.render_combined_row("Сведения о примененных мерах обеспечения производства по материалам о грубом дисциплинарном проступке: ", "")
		self.render_combined_row("Иные сведения: ", "")
		self.render_combined_row("К протоколу прилагаются: ", "материалы служебного разбирательства")
		self.add_empty_paragraphs(1)
		self.render_combined_row("Подпись военнослужащего: ", f"нет, в связи с отсутствием {pers_name_formatted}")
		self.add_paragraph("или отметка об отказе", small_underline)
		self.render_combined_row("Подпись лица составляющего протокол: ", "")
		self.render_combined_row("Копию протокола получил: ",  f"нет, в связи с отсутствием {pers_name_formatted}")
		self.add_paragraph("подпись военнослужащего в отношении которого составлен протокол", small_underline)

		self.render_date_placeholder()

		settings = PersFullNameSettings(1, False, True, True, False, True, False)
		soldier_full_str = self.get_person_full_str(settings)
		self.add_paragraph(f"Решение командира войсковой части {self.get_military_unit()}:", self.align_justify_settings)
		resolution_text = f"За грубый дисциплинарный проступок отсутствие в месте военной службы без уважительных причин более 4 (четырех) часов подряд в течение установленного ежедневного служебного времени, {soldier_full_str}, ПРЕДУПРЕДИТЬ О НЕПОЛНОМ СЛУЖЕБНОМ СООТВЕТСТВИИ."

		paragraph_settings = ParagraphSettings()
		paragraph_settings.align_justify = True
		paragraph_settings.first_line_indent = Mm(12.5)
		paragraph_settings.is_underline = True
		self.add_paragraph(resolution_text, paragraph_settings)

		self.render_date_placeholder()

		commander4_info = self.get_commander_generic("commander_4_level", "КОМАНДИР ВОЙСКОВОЙ ЧАСТИ")
		self.add_paragraph(f"{commander4_info['position'].upper()} {self.get_military_unit()}", self.bold_center_settings)
		self.add_paragraph(commander4_info['rank'], self.bold_center_settings)
		self.add_paragraph(commander4_info['name'], self.bold_right_settings)

		super().render()

	def render_combined_row(self, left_part, right_part):
		p = self.word_document.add_paragraph(left_part)
		p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
		runner = p.add_run(right_part)
		runner.underline = True
		p.add_run("_________")

	def render_date_placeholder(self):
		self.add_empty_paragraphs(1)
		self.add_paragraph("«     » _________ 2023 г.", self.align_justify_settings)
		self.add_empty_paragraphs(1)

	def add_paragraph_with_underline(self, text):
		p = self.add_paragraph(text, self.align_center_underline)
		p.add_run("________________")

