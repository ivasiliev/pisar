from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from classes.document_in_report import DocumentInReport
from classes.paragraph_settings import ParagraphSettings


class DocDiarySoldierInfo(DocumentInReport):
	def get_name(self):
		return "Индивидуальные данные военнослужащего"

	def get_name_for_file(self):
		return f"{self.get_name()} ({self.get_soldier_info().full_name}).docx"

	def render(self):
		s_info = self.get_soldier_info()
		rep_settings = self.get_report_settings()

		f_12_centered = ParagraphSettings()
		f_12_centered.align_center = True
		f_12_centered.font_size = Pt(12)

		f_12_bold_centered = ParagraphSettings()
		f_12_bold_centered.align_center = True
		f_12_bold_centered.font_size = Pt(12)
		f_12_bold_centered.is_bold = True

		f_10_bold_centered = ParagraphSettings()
		f_10_bold_centered.align_center = True
		f_10_bold_centered.font_size = Pt(10)
		f_10_bold_centered.is_bold = True

		f_10 = ParagraphSettings()
		f_10.font_size = Pt(10)

		self.add_paragraph("4. Индивидуальные данные военнослужащего", self.bold_center_settings)
		self.add_paragraph("(оформляются на каждого военнослужащего до проведения индивидуальных бесед из личных дел и других документов, уточняются после беседы)", f_12_centered)
		self.add_empty_paragraphs(1)
		# TODO try to do underline
		self.add_paragraph(f"Фамилия, имя и отчество: {s_info.full_name}", f_10)
		left_part = f"Воинское звание: {self.get_person_rank(s_info.rank, 0)}"
		right_part = f"Дата рождения: {s_info.get_dob()}"
		self.add_paragraph_left_right(left_part, right_part, Pt(10), WD_TABLE_ALIGNMENT.LEFT)
		self.add_paragraph(f"Назначен в подразделение, должность: {self.get_soldier_address(0)}; {s_info.position}.", f_10)
		self.add_empty_paragraphs(1)

		self.add_paragraph("Внешние приметы", f_10_bold_centered)
		captions = ["", ""]
		table_settings = {"ps": None, "cols_width": [30, 150], "alignment": WD_PARAGRAPH_ALIGNMENT.LEFT, "font_size": Pt(10)}
		# TODO fill data
		row_data = [
			["Рост:", str(rep_settings["height"])]
			, ["Вес:", str(rep_settings["weight"])]
			, ["Особые приметы:", rep_settings["signs"]]
			, ["Татуировки:", rep_settings["tatoo"]]
			, ["Привычки:", rep_settings["habits"]]
		]
		self.add_table(captions, row_data, table_settings)
		self.add_empty_paragraphs(1)

		self.add_paragraph("Сведения о родителях, родственниках", f_10_bold_centered)
		self.add_paragraph("(ФИО, адрес, дата рождения, место работы, номер телефона, дополнительные сведения)", f_10_bold_centered)
		row_data = [
			["Отец:", rep_settings["father_name"]]
			, ["Мать:", rep_settings["mother_name"]]
			, ["Брат:", rep_settings["siblings"]]
			, ["Сестра:", ""]
			, ["Другие ближайшие родственники:", ""]
			, ["Жена (Девушка):", rep_settings["spouse"]]
			, ["Прочие:", ""]
		]
		self.add_table(captions, row_data, table_settings)
		self.word_document.add_page_break()

		super().render()