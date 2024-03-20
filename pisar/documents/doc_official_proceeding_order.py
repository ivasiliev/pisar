from docx.shared import Mm
from helpers.text_helper import capitalize_first_letter

from classes.document_in_report import DocumentInReport
from classes.paragraph_settings import ParagraphSettings
from classes.pers_full_name_settings import PersFullNameSettings


class DocOfficialProceedingOrder(DocumentInReport):
	def get_name(self):
		return "Приказ командира войсковой части"

	def get_name_for_file(self):
		return f"{self.get_name()} ({self.get_soldier_info().full_name}).docx"

	def render(self, custom_margins=None):
		s_info = self.get_soldier_info()
		rep_settings = self.get_report_settings()
		date_of_event = self.get_date_format_1(rep_settings["date_of_event"])

		paragraph_settings = ParagraphSettings()
		paragraph_settings.align_right = True
		paragraph_settings.right_indent = Mm(7.5)
		self.add_paragraph("Для служебного пользования", paragraph_settings)
		self.add_paragraph("(п. 648 Перечня сведений ВС)", paragraph_settings)
		paragraph_settings.right_indent = Mm(30)
		self.add_paragraph("Экз. №____", paragraph_settings)
		self.add_empty_paragraphs(1)

		self.add_paragraph("П Р И К А З", self.bold_title)
		self.add_paragraph(f"КОМАНДИРА ВОЙСКОВОЙ ЧАСТИ {self.get_military_unit()}", self.bold_center_settings)
		self.add_paragraph("«___» _________ 2024 г. № ___", self.align_center_settings)
		self.add_paragraph("г. Донецк", self.align_center_settings)
		self.add_empty_paragraphs(1)

		self.add_paragraph("По факту грубого дисциплинарного проступка", self.bold_center_settings)
		paragraph_settings = ParagraphSettings()
		paragraph_settings.is_bold = True
		paragraph_settings.is_underline = True
		paragraph_settings.align_center = True

		# TODO this can be in base class
		rnk = self.get_person_rank(s_info.rank, 2)
		nm = self.get_person_name_declension(s_info.full_name, 2)
		self.add_paragraph(f"{rnk} {nm}", paragraph_settings)
		self.add_empty_paragraphs(2)
		
		# custom settings
		line_spacing = 1.0
		custom_ident_align_justify_settings = self.ident_align_justify_settings
		custom_ident_align_justify_settings.line_spacing = line_spacing

		txt =f"Несмотря на меры, принимаемые командованием войсковой части {self.get_military_unit()} по профилактике правонарушений, связанных с уклонением военнослужащих от исполнения служебных обязанностей и прохождения военной службы, продолжаются факты самовольного оставления воинской части или места несения службы военнослужащими войсковой части {self.get_military_unit()}."
		self.add_paragraph(txt, custom_ident_align_justify_settings)

		settings = PersFullNameSettings(0, False, False, True, False, True, False)
		sold_str = self.get_person_full_str(settings)
		txt = f"В ходе проведения служебного разбирательства было установлено, что {date_of_event} {sold_str} отсутствовал на месте несения воинской службы без уважительных причин более 4 (четырех) часов подряд в течение установленного ежедневного служебного времени, тем самым совершил грубый дисциплинарный проступок."
		self.add_paragraph(txt, custom_ident_align_justify_settings)

		rnk = self.get_person_rank(s_info.rank, 1)
		nm = self.get_person_name_declension(s_info.full_name, 1)
		txt = f"Проведенные розыскные мероприятия, опрос сослуживцев, поиск в лечебных учреждениях, военных комендатурах, а также отделениях полиции {rnk} {nm} результата не дали, установить причины отсутствия военнослужащего, а также его местонахождение не удалось."
		self.add_paragraph(txt, custom_ident_align_justify_settings)

		self.add_paragraph("Причинами данного правонарушения явились:", custom_ident_align_justify_settings)

		commander_company_text = self.get_commander_company_full_str(2, False)
		txt = "невыполнение требований статьи 144, 145 Устава Внутренней Службы Вооруженных Сил Российской Федерации в " \
		      "части, касающейся воспитания, поддержания воинской дисциплины, морально-психологического состояния " \
		      "личного состава в роте, а также в части касающееся знаний деловых и морально-психологических качеств и особенностей всех военнослужащих роты, постоянного " \
		            f"проведения с ними индивидуальной работы по воинскому воспитанию {commander_company_text};"

		self.add_paragraph(txt, custom_ident_align_justify_settings)

		commander_platoon = self.get_commander_platoon_full_str(1)
		txt = "невыполнение требований статьи 152, 153 Устава Внутренней Службы Вооруженных Сил Российской Федерации в " \
		      "части, касающейся воспитания, поддержания воинской дисциплины, морально–психологического состояния во " \
		      f"взводе со стороны {commander_platoon};"
		self.add_paragraph(txt, custom_ident_align_justify_settings)

		commander_squad = self.get_commander_generic_full_str("commander_squad", 2, False)
		txt = f"невыполнение требований статей 156, 157 Устава Внутренней Службы Вооружённых Сил Российской Федерации, в части касающейся воспитания, поддержания воинской дисциплины, укрепления морально-политического и психологического состояния подчиненного личного состава взводе {commander_squad};"
		self.add_paragraph(txt, custom_ident_align_justify_settings)

		settings = PersFullNameSettings(2, False, False, True, False, True, False)
		sold_str = self.get_person_full_str(settings)
		txt = f"невыполнение требований статьи 160, 161 Устава Внутренней Службы Вооруженных Сил Российской Федерации в части, касающейся точного и своевременного исполнения возложенных на него обязанностей, поставленных задач {sold_str}."
		self.add_paragraph(txt, custom_ident_align_justify_settings)

		rnk = self.get_person_rank(s_info.rank, 0)
		nm = self.get_person_name_declension(s_info.full_name, 0)
		self.add_paragraph(
			f"Исходя из материала служебного разбирательства следует, что {rnk} {nm} отсутствовал на месте несения воинской службы без уважительных причин более 4 (четырех) часов подряд в течение установленного ежедневного служебного времени, что в соответствии с Приложением 7 к Дисциплинарному Уставу Вооруженных Сил Российской Федерации является грубым дисциплинарным проступком."
			, custom_ident_align_justify_settings)
		p1 = self.add_paragraph("На основании вышеизложенного, ", custom_ident_align_justify_settings)
		runner = p1.add_run("ПРИКАЗЫВАЮ:")
		runner.bold = True

		commander_company_text = self.get_commander_company_full_str(3, False)
		txt = ("1. За невыполнение требований статей 144,145 Устава внутренней службы Вооруженных Сил Российской "
		       "Федерации, в части касающейся организации и проведения им работы по повседневному воспитанию, "
		       "поддержанию воинской дисциплины, укреплению морально-политического и психологического состояния "
		       "подчиненного личного состава"
		       f" {commander_company_text} провести дополнительные занятия в роте по "
		       "ознакомлению со статьями Уголовного кодекса Российской Федерации и наказаниями за их нарушение.")
		self.add_paragraph(txt, custom_ident_align_justify_settings)

		commander_platoon = self.get_commander_platoon_full_str(3)
		txt = (f"2. За невыполнение требований статей 152, 153 Устава внутренней службы Вооруженных Сил Российской "
		       f"Федерации, в части касающейся воспитания, поддержания воинской дисциплины, укрепления "
		       f"морально-политического и психологического состояния солдат подчиненного взвода,"
		       f" {commander_platoon}, строго указать на исполнение служебных и должностных обязанностей.")
		self.add_paragraph(txt, custom_ident_align_justify_settings)

		commander_squad = self.get_commander_generic_full_str("commander_squad", 3, False)
		txt = (f"3. За невыполнение требований статей 156, 157 Устава Внутренней Службы Вооружённых Сил Российской "
		       f"Федерации, в части касающейся воспитания, поддержания воинской дисциплины, укрепления "
		       f"морально-политического и психологического состояния подчиненного личного состава взвода "
		       f"{commander_squad}, строго указать на низкую дисциплину в отделении.")
		self.add_paragraph(txt, custom_ident_align_justify_settings)

		settings = PersFullNameSettings(3, False, False, True, False, True, False)
		sold_str = capitalize_first_letter(self.get_person_full_str(settings))
		txt = (f"4. {sold_str} за грубый дисциплинарный проступок, отсутствие на месте несения службы без уважительных "
		       f"причин более 4 (четырех) часов подряд в течение установленного ежедневного служебного времени, "
		       f"объявить СТРОГИЙ ВЫГОВОР.")
		self.add_paragraph(txt, custom_ident_align_justify_settings)

		self.add_paragraph(f"5. Контроль за исполнением настоящего приказа возложить на начальника штаба - заместителя командира войсковой части {self.get_military_unit()}.", custom_ident_align_justify_settings)

		self.add_paragraph(f"6. Приказ довести до личного состава в части касающейся.", custom_ident_align_justify_settings)

		self.add_empty_paragraphs(3)

		# TODO this code is SIMILAR in Performance_Characteristics

		line_spacing = 0.96
		par_set_center = ParagraphSettings()
		par_set_center.is_bold = True
		par_set_center.align_center = True
		par_set_center.line_spacing = line_spacing

		par_set_right = ParagraphSettings()
		par_set_right.align_right = True
		par_set_right.is_bold = True
		par_set_right.line_spacing = line_spacing

		self.add_commander(rep_settings["commander_4_level"], self.get_military_unit(), par_set_center, par_set_right)
		self.add_empty_paragraphs(2)
		self.add_commander(rep_settings["commander_3_level"], self.get_military_unit(), par_set_center, par_set_right)

		super().render()

	def add_commander(self, commander_info, military_unit, par_set_center, par_set_right):
		c_name = self.get_person_name_short_format_1(commander_info["name"])
		c_rank = commander_info["rank"]
		if self.get_report_settings()["is_guard"]:
			c_rank = "гвардии " + c_rank
		c_position = commander_info["position"] + " " + military_unit
		self.add_paragraph(c_position.upper(), self.bold_center_settings)
		self.add_paragraph(c_rank, par_set_center)
		self.add_paragraph(c_name, par_set_right)
